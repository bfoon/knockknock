"""
Knock-Knock export module.

Generates branded Word (.docx) and Excel (.xlsx) reports from a Questionnaire
+ LiveSession, with embedded charts that mirror what was shown on stage.

Drop this file in `polls/exports.py` (or `core/exports.py` if you also want it
for games). It is framework-agnostic Python — the Django views just call
build_word_report(...) / build_excel_report(...) and stream the bytes back.

Dependencies (add to requirements.txt):
    python-docx>=1.1
    openpyxl>=3.1
    matplotlib>=3.8
"""

from __future__ import annotations

import io
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Iterable

import matplotlib

matplotlib.use("Agg")  # headless rendering for Django workers
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Inches

from openpyxl import Workbook
from openpyxl.chart import BarChart, PieChart, LineChart, DoughnutChart, Reference
from openpyxl.chart.label import DataLabelList
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage


# ─────────────────────────────────────────────────────────────────────────────
# Brand palette — matches the in-app "midnight" template feel
# ─────────────────────────────────────────────────────────────────────────────

BRAND = {
    "ink":      "#0B1020",   # near-black background colour used in headings
    "accent":   "#7C5CFF",   # primary purple
    "accent2":  "#22D3EE",   # cyan secondary
    "warm":     "#FB7185",   # pink/red used for highlights
    "lime":     "#A3E635",
    "amber":    "#FBBF24",
    "muted":    "#94A3B8",
    "paper":    "#F8FAFC",
    "rule":     "#E2E8F0",
}

# colour wheel used when a question has many options — cycled in order
CHART_PALETTE = [
    BRAND["accent"], BRAND["accent2"], BRAND["warm"],
    BRAND["lime"], BRAND["amber"], "#F472B6", "#34D399", "#60A5FA",
]


# ─────────────────────────────────────────────────────────────────────────────
# Lightweight data containers — what the views pass in
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class QuestionResult:
    text: str
    qtype: str            # "mcq" | "word" | "scale" | "open" | "ranking"
    chart_type: str       # "bar" | "donut" | "pie" | "line" | "wordcloud" | …
    options: list[str] = field(default_factory=list)     # MCQ / ranking labels
    counts:  list[int] = field(default_factory=list)     # tally per option
    words:   list[str] = field(default_factory=list)     # for word-cloud
    scale_values: list[float] = field(default_factory=list)
    open_answers: list[str]   = field(default_factory=list)

    # ── derived stats ──────────────────────────────────────────────────────
    @property
    def total_responses(self) -> int:
        if self.qtype == "mcq" or self.qtype == "ranking":
            return sum(self.counts)
        if self.qtype == "word":
            return len(self.words)
        if self.qtype == "scale":
            return len(self.scale_values)
        if self.qtype == "open":
            return len(self.open_answers)
        return 0

    @property
    def average(self) -> float | None:
        if self.qtype == "scale" and self.scale_values:
            return sum(self.scale_values) / len(self.scale_values)
        return None

    @property
    def top_choice(self) -> str | None:
        if not self.counts or not self.options:
            return None
        i = max(range(len(self.counts)), key=lambda k: self.counts[k])
        return self.options[i]


@dataclass
class ReportData:
    title: str
    description: str
    owner_name: str
    session_code: str
    mode: str
    started_at: datetime
    ended_at: datetime | None
    participant_count: int
    questions: list[QuestionResult]


# ─────────────────────────────────────────────────────────────────────────────
# Matplotlib chart rendering  →  PNG bytes (used for the Word report)
# ─────────────────────────────────────────────────────────────────────────────

def _style_axes(ax, *, light: bool = True) -> None:
    """Knock-Knock house style for matplotlib axes."""
    fg = "#1E293B" if light else BRAND["paper"]
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(BRAND["rule"])
    ax.spines["bottom"].set_color(BRAND["rule"])
    ax.tick_params(colors=fg, labelsize=10)
    ax.yaxis.label.set_color(fg)
    ax.xaxis.label.set_color(fg)
    ax.grid(axis="y", color=BRAND["rule"], linewidth=0.6, alpha=0.7)
    ax.set_axisbelow(True)


def _new_figure(figsize=(8.5, 4.6)):
    fig, ax = plt.subplots(figsize=figsize, dpi=160)
    fig.patch.set_facecolor("white")
    return fig, ax


def _render_bar(q: QuestionResult) -> bytes:
    fig, ax = _new_figure()
    colors = [CHART_PALETTE[i % len(CHART_PALETTE)] for i in range(len(q.options))]
    bars = ax.bar(q.options, q.counts, color=colors, width=0.62,
                  edgecolor="white", linewidth=2)
    # rounded look via clipping isn't trivial — fake it with a soft top label
    for bar, v in zip(bars, q.counts):
        ax.text(bar.get_x() + bar.get_width() / 2, v + max(q.counts) * 0.02,
                str(v), ha="center", va="bottom",
                fontsize=11, fontweight="bold", color="#0F172A")
    ax.set_ylabel("Responses")
    _style_axes(ax)
    plt.xticks(rotation=0, ha="center")
    fig.tight_layout()
    return _fig_to_png(fig)


def _render_pie_or_donut(q: QuestionResult, donut: bool) -> bytes:
    fig, ax = _new_figure(figsize=(7.5, 4.6))
    colors = [CHART_PALETTE[i % len(CHART_PALETTE)] for i in range(len(q.options))]
    wedges, _texts, autotexts = ax.pie(
        q.counts, labels=None, colors=colors, autopct="%1.0f%%",
        startangle=90, pctdistance=0.78 if donut else 0.62,
        wedgeprops=dict(width=0.42 if donut else 1.0, edgecolor="white", linewidth=2),
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontweight("bold")
        t.set_fontsize(10)
    ax.legend(wedges, q.options, loc="center left",
              bbox_to_anchor=(1.0, 0.5), frameon=False, fontsize=10)
    ax.set_aspect("equal")
    fig.tight_layout()
    return _fig_to_png(fig)


def _render_line(q: QuestionResult) -> bytes:
    fig, ax = _new_figure()
    ax.plot(q.options, q.counts, marker="o", linewidth=3,
            color=BRAND["accent"], markerfacecolor=BRAND["accent2"],
            markersize=10, markeredgecolor="white", markeredgewidth=2)
    ax.fill_between(range(len(q.options)), q.counts, alpha=0.12,
                    color=BRAND["accent"])
    for x, v in enumerate(q.counts):
        ax.text(x, v + max(q.counts) * 0.04, str(v),
                ha="center", fontsize=10, fontweight="bold")
    ax.set_ylabel("Responses")
    _style_axes(ax)
    fig.tight_layout()
    return _fig_to_png(fig)


def _render_scale(q: QuestionResult) -> bytes:
    fig, ax = _new_figure()
    # histogram 1..10
    bins = list(range(1, 12))
    ax.hist(q.scale_values, bins=bins, color=BRAND["accent"],
            edgecolor="white", linewidth=2, align="left", rwidth=0.85)
    if q.average is not None:
        ax.axvline(q.average, color=BRAND["warm"], linewidth=2.5,
                   linestyle="--", label=f"Avg {q.average:.2f}")
        ax.legend(frameon=False)
    ax.set_xticks(range(1, 11))
    ax.set_xlabel("Score")
    ax.set_ylabel("Responses")
    _style_axes(ax)
    fig.tight_layout()
    return _fig_to_png(fig)


def _render_wordcloud(q: QuestionResult) -> bytes:
    """Custom 'tag cloud' without the wordcloud dep — sized by frequency."""
    fig, ax = _new_figure(figsize=(8.5, 4.6))
    ax.set_xlim(0, 100); ax.set_ylim(0, 100); ax.axis("off")
    freqs = Counter(w.lower().strip() for w in q.words if w.strip())
    if not freqs:
        ax.text(50, 50, "(no responses)", ha="center", va="center",
                fontsize=14, color=BRAND["muted"])
        return _fig_to_png(fig)
    max_f = max(freqs.values())
    # simple placement: alternating columns/rows by rank
    items = freqs.most_common(40)
    import random
    random.seed(7)
    placed = []
    for word, f in items:
        size = 14 + (f / max_f) * 38
        color = CHART_PALETTE[hash(word) % len(CHART_PALETTE)]
        # try a handful of spots
        for _ in range(50):
            x, y = random.uniform(10, 90), random.uniform(15, 85)
            ok = all(abs(x - px) > 11 or abs(y - py) > 6 for px, py, _ in placed)
            if ok:
                placed.append((x, y, size))
                ax.text(x, y, word, ha="center", va="center",
                        fontsize=size, color=color, fontweight="bold")
                break
    fig.tight_layout()
    return _fig_to_png(fig)


def _render_ranking(q: QuestionResult) -> bytes:
    fig, ax = _new_figure(figsize=(8.5, max(3.5, len(q.options) * 0.55)))
    # horizontal bars, descending
    order = sorted(range(len(q.options)),
                   key=lambda i: q.counts[i], reverse=True)
    labels = [q.options[i] for i in order][::-1]
    values = [q.counts[i] for i in order][::-1]
    colors = [CHART_PALETTE[i % len(CHART_PALETTE)] for i in range(len(labels))]
    bars = ax.barh(labels, values, color=colors, edgecolor="white", linewidth=2)
    for bar, v in zip(bars, values):
        ax.text(v + max(values) * 0.02, bar.get_y() + bar.get_height() / 2,
                str(v), va="center", fontsize=11, fontweight="bold")
    ax.set_xlabel("Points")
    _style_axes(ax)
    fig.tight_layout()
    return _fig_to_png(fig)


def _fig_to_png(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return buf.getvalue()


def render_chart_png(q: QuestionResult) -> bytes:
    """Dispatch on chart_type / qtype → return PNG bytes."""
    if q.qtype == "scale":
        return _render_scale(q)
    if q.qtype == "word":
        return _render_wordcloud(q)
    if q.qtype == "ranking":
        return _render_ranking(q)
    if q.qtype == "open":
        # no chart for open text — caller will render the verbatim list
        return b""
    # MCQ
    if q.chart_type in ("pie",):
        return _render_pie_or_donut(q, donut=False)
    if q.chart_type in ("donut", "doughnut"):
        return _render_pie_or_donut(q, donut=True)
    if q.chart_type == "line":
        return _render_line(q)
    return _render_bar(q)


# ─────────────────────────────────────────────────────────────────────────────
# WORD report
# ─────────────────────────────────────────────────────────────────────────────

def _hex(c: str) -> RGBColor:
    return RGBColor.from_string(c.lstrip("#"))


def _shade_cell(cell, hex_color: str) -> None:
    """Fill a table cell with a solid colour (python-docx has no helper)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _set_cell_border(cell, **kwargs):
    """e.g. _set_cell_border(cell, bottom={'sz':6, 'color':'7C5CFF'})"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, attrs in kwargs.items():
        e = tc_borders.find(qn(f"w:{edge}"))
        if e is None:
            e = OxmlElement(f"w:{edge}")
            tc_borders.append(e)
        e.set(qn("w:val"), attrs.get("val", "single"))
        e.set(qn("w:sz"), str(attrs.get("sz", 4)))
        e.set(qn("w:color"), attrs.get("color", "auto"))


def _add_pill(p, text: str, fill: str, fg: str = "FFFFFF") -> None:
    """Faked rounded chip via shaded inline text in its own little run."""
    run = p.add_run(f"  {text}  ")
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = _hex(fg)
    # shade behind text
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    rPr.append(shd)


def _add_horizontal_rule(doc, color: str = BRAND["rule"]) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.lstrip("#"))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _question_type_label(qtype: str) -> str:
    return {
        "mcq": "Multiple Choice",
        "word": "Word Cloud",
        "scale": "Scale (1-10)",
        "open": "Open Text",
        "ranking": "Ranking",
    }.get(qtype, qtype)


def build_word_report(data: ReportData) -> bytes:
    """Return .docx bytes for download."""
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── Global default font ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── COVER ─────────────────────────────────────────────────────────────
    # Brand strap
    strap = doc.add_paragraph()
    r = strap.add_run("KNOCK-KNOCK")
    r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = _hex(BRAND["accent"])
    strap.add_run("    ·    LIVE SESSION RESULTS").font.color.rgb = _hex(BRAND["muted"])
    strap.runs[1].font.size = Pt(10)
    strap.runs[1].font.bold = True

    # Title
    title = doc.add_paragraph()
    tr = title.add_run(data.title)
    tr.font.size = Pt(32); tr.font.bold = True
    tr.font.color.rgb = _hex(BRAND["ink"])
    title.paragraph_format.space_after = Pt(4)

    if data.description:
        sub = doc.add_paragraph()
        sr = sub.add_run(data.description)
        sr.font.size = Pt(12); sr.font.color.rgb = _hex(BRAND["muted"])
        sub.paragraph_format.space_after = Pt(10)

    # Meta pills
    pills = doc.add_paragraph()
    _add_pill(pills, f"Code  {data.session_code}", BRAND["ink"])
    _add_pill(pills, f"Mode  {data.mode.capitalize()}", BRAND["accent"])
    _add_pill(pills, f"{data.participant_count} participants", BRAND["accent2"], fg="0B1020")
    _add_pill(pills, f"{len(data.questions)} questions", BRAND["warm"])
    pills.paragraph_format.space_after = Pt(6)

    # Stamp
    stamp = doc.add_paragraph()
    s = stamp.add_run(
        f"Hosted by {data.owner_name}  ·  "
        f"{data.started_at.strftime('%b %d, %Y · %H:%M')}"
        + (f" – {data.ended_at.strftime('%H:%M')}" if data.ended_at else "")
    )
    s.font.size = Pt(9); s.font.color.rgb = _hex(BRAND["muted"])

    _add_horizontal_rule(doc, BRAND["accent"])

    # ── EXECUTIVE SUMMARY ────────────────────────────────────────────────
    h = doc.add_paragraph()
    hr = h.add_run("Executive Summary")
    hr.font.size = Pt(18); hr.font.bold = True
    hr.font.color.rgb = _hex(BRAND["ink"])
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(6)

    summary = doc.add_paragraph()
    summary.add_run(
        f"This report summarises {len(data.questions)} question"
        f"{'s' if len(data.questions) != 1 else ''} answered live by "
        f"{data.participant_count} participants. Each section below shows "
        "the question, its chart, and the underlying response data."
    ).font.size = Pt(11)
    summary.paragraph_format.space_after = Pt(8)

    # KPI strip — 3 cards as a 1-row table
    kpi_tbl = doc.add_table(rows=1, cols=3)
    kpi_tbl.autofit = False
    for cell, (label, value, fill) in zip(
        kpi_tbl.rows[0].cells,
        [
            ("PARTICIPANTS", str(data.participant_count), BRAND["accent"]),
            ("QUESTIONS",   str(len(data.questions)),     BRAND["accent2"]),
            ("RESPONSES",   str(sum(q.total_responses for q in data.questions)),
                                                          BRAND["warm"]),
        ],
    ):
        _shade_cell(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # label
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(label)
        r1.font.size = Pt(8); r1.font.bold = True
        r1.font.color.rgb = _hex("FFFFFF")
        # value
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(value)
        r2.font.size = Pt(26); r2.font.bold = True
        r2.font.color.rgb = _hex("FFFFFF")
        # spacer
        cell.add_paragraph()
    # equalise column widths
    for col in kpi_tbl.columns:
        for c in col.cells:
            c.width = Cm(5.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── TABLE OF CONTENTS (simple list) ──────────────────────────────────
    h2 = doc.add_paragraph()
    h2r = h2.add_run("Contents")
    h2r.font.size = Pt(14); h2r.font.bold = True
    h2r.font.color.rgb = _hex(BRAND["ink"])
    for i, q in enumerate(data.questions, 1):
        toc = doc.add_paragraph()
        n = toc.add_run(f"  {i:02d}    ")
        n.font.color.rgb = _hex(BRAND["accent"])
        n.font.bold = True
        t = toc.add_run(q.text)
        t.font.color.rgb = _hex(BRAND["ink"])
        meta = toc.add_run(f"   ·  {_question_type_label(q.qtype)}")
        meta.font.color.rgb = _hex(BRAND["muted"])
        meta.font.size = Pt(9)

    doc.add_page_break()

    # ── PER-QUESTION SECTIONS ────────────────────────────────────────────
    for i, q in enumerate(data.questions, 1):
        # Numbered header
        head = doc.add_paragraph()
        num = head.add_run(f"Q{i:02d}  ")
        num.font.size = Pt(11); num.font.bold = True
        num.font.color.rgb = _hex(BRAND["accent"])
        ttl = head.add_run(q.text)
        ttl.font.size = Pt(17); ttl.font.bold = True
        ttl.font.color.rgb = _hex(BRAND["ink"])

        # Tag line
        tag = doc.add_paragraph()
        _add_pill(tag, _question_type_label(q.qtype), BRAND["ink"])
        _add_pill(tag, q.chart_type.upper(), BRAND["accent"])
        _add_pill(tag, f"{q.total_responses} responses", BRAND["accent2"], fg="0B1020")
        if q.average is not None:
            _add_pill(tag, f"Avg {q.average:.2f}", BRAND["warm"])
        tag.paragraph_format.space_after = Pt(4)

        # Chart image (skip for open text)
        png = render_chart_png(q)
        if png:
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_p.add_run()
            run.add_picture(io.BytesIO(png), width=Cm(16.5))

        # Data table
        if q.qtype in ("mcq", "ranking") and q.options:
            total = max(sum(q.counts), 1)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            for cell, label in zip(hdr, ("Option", "Count", "Share")):
                _shade_cell(cell, BRAND["ink"])
                p = cell.paragraphs[0]
                r = p.add_run(label)
                r.font.bold = True; r.font.color.rgb = _hex("FFFFFF")
                r.font.size = Pt(10)
            for opt, cnt in zip(q.options, q.counts):
                row = tbl.add_row().cells
                row[0].text = opt
                row[1].text = str(cnt)
                row[2].text = f"{cnt / total * 100:.1f}%"
            # widths
            for r_ in tbl.rows:
                r_.cells[0].width = Cm(10)
                r_.cells[1].width = Cm(3)
                r_.cells[2].width = Cm(3)

        elif q.qtype == "scale" and q.scale_values:
            tbl = doc.add_table(rows=2, cols=4)
            tbl.style = "Light Grid Accent 1"
            for cell, label in zip(tbl.rows[0].cells,
                                   ("Average", "Median", "Min", "Max")):
                _shade_cell(cell, BRAND["ink"])
                r = cell.paragraphs[0].add_run(label)
                r.font.bold = True; r.font.color.rgb = _hex("FFFFFF")
                r.font.size = Pt(10)
            vs = sorted(q.scale_values)
            median = vs[len(vs) // 2]
            for cell, val in zip(tbl.rows[1].cells, (
                f"{q.average:.2f}", f"{median:g}", f"{min(vs):g}", f"{max(vs):g}"
            )):
                cell.paragraphs[0].add_run(val).font.size = Pt(11)

        elif q.qtype == "open" and q.open_answers:
            for ans in q.open_answers[:40]:    # cap to keep doc sensible
                p = doc.add_paragraph()
                bullet = p.add_run("▸  ")
                bullet.font.color.rgb = _hex(BRAND["accent"])
                bullet.font.bold = True
                p.add_run(ans)
            if len(q.open_answers) > 40:
                p = doc.add_paragraph()
                more = p.add_run(f"…and {len(q.open_answers) - 40} more responses "
                                 f"(see Excel export for full list)")
                more.font.italic = True
                more.font.color.rgb = _hex(BRAND["muted"])

        elif q.qtype == "word" and q.words:
            top = Counter(w.lower() for w in q.words if w.strip()).most_common(15)
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Light Grid Accent 1"
            for cell, label in zip(tbl.rows[0].cells, ("Word", "Mentions")):
                _shade_cell(cell, BRAND["ink"])
                r = cell.paragraphs[0].add_run(label)
                r.font.bold = True; r.font.color.rgb = _hex("FFFFFF")
                r.font.size = Pt(10)
            for w, cnt in top:
                row = tbl.add_row().cells
                row[0].text = w
                row[1].text = str(cnt)

        # spacer + separator (skip after last)
        if i < len(data.questions):
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            _add_horizontal_rule(doc)

    # ── Footer (every page) ──────────────────────────────────────────────
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"Knock-Knock  ·  {data.title}  ·  {data.session_code}")
    fr.font.size = Pt(8); fr.font.color.rgb = _hex(BRAND["muted"])

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# EXCEL report — uses native Excel charts (open in Excel & they're live)
# ─────────────────────────────────────────────────────────────────────────────

# colour helpers (openpyxl wants RGB w/o #)
def _xhex(c: str) -> str:
    return c.lstrip("#").upper()

THIN  = Side(border_style="thin", color=_xhex(BRAND["rule"]))
ACCENT_BOTTOM = Side(border_style="medium", color=_xhex(BRAND["accent"]))


def _style_header_row(ws, row: int, cols: int) -> None:
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=_xhex(BRAND["ink"]))
        cell.alignment = Alignment(horizontal="left", vertical="center")
        cell.border = Border(top=THIN, bottom=ACCENT_BOTTOM,
                             left=THIN, right=THIN)
    ws.row_dimensions[row].height = 24


def _write_title_banner(ws, text: str, sub: str = "") -> int:
    """Draws a coloured title banner at the top — returns the next row."""
    ws["A1"] = text
    ws["A1"].font = Font(name="Calibri", size=20, bold=True, color="FFFFFF")
    ws["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws["A1"].fill = PatternFill("solid", fgColor=_xhex(BRAND["ink"]))
    ws.merge_cells("A1:F1")
    ws.row_dimensions[1].height = 38

    if sub:
        ws["A2"] = sub
        ws["A2"].font = Font(name="Calibri", size=11, italic=True,
                             color=_xhex(BRAND["muted"]))
        ws["A2"].alignment = Alignment(horizontal="left", indent=1)
        ws.merge_cells("A2:F2")
        return 4
    return 3


def _add_summary_sheet(wb, data: ReportData) -> None:
    ws = wb.active
    ws.title = "Summary"
    ws.sheet_view.showGridLines = False

    row = _write_title_banner(
        ws,
        data.title,
        f"Session {data.session_code}  ·  Hosted by {data.owner_name}  ·  "
        f"{data.started_at.strftime('%b %d, %Y · %H:%M')}",
    )

    # KPI strip
    kpis = [
        ("Participants", data.participant_count, BRAND["accent"]),
        ("Questions",    len(data.questions),    BRAND["accent2"]),
        ("Responses",    sum(q.total_responses for q in data.questions),
                                                 BRAND["warm"]),
        ("Mode",         data.mode.capitalize(), BRAND["lime"]),
    ]
    for i, (label, value, color) in enumerate(kpis):
        col = 1 + i * 2          # A, C, E, G  — leave gaps
        end_col = col + 1
        # label row
        c_label = ws.cell(row=row, column=col, value=label.upper())
        c_label.font = Font(bold=True, size=9, color="FFFFFF", name="Calibri")
        c_label.fill = PatternFill("solid", fgColor=_xhex(color))
        c_label.alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells(start_row=row, start_column=col,
                       end_row=row, end_column=end_col)
        # value row
        c_val = ws.cell(row=row + 1, column=col, value=value)
        c_val.font = Font(bold=True, size=22, color=_xhex(BRAND["ink"]),
                          name="Calibri")
        c_val.fill = PatternFill("solid", fgColor="FFFFFF")
        c_val.alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells(start_row=row + 1, start_column=col,
                       end_row=row + 2, end_column=end_col)
        # border around value cells
        for rr in (row + 1, row + 2):
            for cc in (col, end_col):
                cell = ws.cell(row=rr, column=cc)
                cell.border = Border(
                    top=Side(style="thin", color=_xhex(color)),
                    bottom=Side(style="thin", color=_xhex(color)),
                    left=Side(style="thin", color=_xhex(color)),
                    right=Side(style="thin", color=_xhex(color)),
                )

    row += 5

    # Question index table
    ws.cell(row=row, column=1, value="QUESTION INDEX").font = Font(
        bold=True, size=12, color=_xhex(BRAND["ink"]), name="Calibri")
    row += 1
    headers = ["#", "Question", "Type", "Chart", "Responses", "Top result"]
    for i, h in enumerate(headers, 1):
        ws.cell(row=row, column=i, value=h)
    _style_header_row(ws, row, len(headers))
    row += 1
    for i, q in enumerate(data.questions, 1):
        ws.cell(row=row, column=1, value=i)
        ws.cell(row=row, column=2, value=q.text)
        ws.cell(row=row, column=3, value=_question_type_label(q.qtype))
        ws.cell(row=row, column=4, value=q.chart_type)
        ws.cell(row=row, column=5, value=q.total_responses)
        top = q.top_choice
        if top is None and q.average is not None:
            top = f"Avg {q.average:.2f}"
        ws.cell(row=row, column=6, value=top or "—")
        # zebra rows
        if i % 2 == 0:
            for c in range(1, len(headers) + 1):
                ws.cell(row=row, column=c).fill = PatternFill(
                    "solid", fgColor="F8FAFC")
        for c in range(1, len(headers) + 1):
            cell = ws.cell(row=row, column=c)
            cell.font = Font(name="Calibri", size=11)
            cell.alignment = Alignment(vertical="center",
                                       horizontal="left" if c != 1 else "center")
            cell.border = Border(bottom=THIN)
        row += 1

    # Column widths
    widths = [6, 52, 18, 14, 13, 28]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w


def _add_question_sheet(wb, idx: int, q: QuestionResult) -> None:
    safe_name = f"Q{idx:02d}"
    ws = wb.create_sheet(title=safe_name)
    ws.sheet_view.showGridLines = False

    row = _write_title_banner(ws, f"Q{idx:02d}  {q.text}",
                              f"{_question_type_label(q.qtype)}  ·  "
                              f"{q.total_responses} responses")

    # ── MCQ / Ranking: data table + native chart ──────────────────────────
    if q.qtype in ("mcq", "ranking") and q.options:
        total = max(sum(q.counts), 1)
        ws.cell(row=row, column=1, value="Option")
        ws.cell(row=row, column=2, value="Count")
        ws.cell(row=row, column=3, value="Share")
        _style_header_row(ws, row, 3)
        data_start = row + 1
        for i, (opt, cnt) in enumerate(zip(q.options, q.counts)):
            r = data_start + i
            ws.cell(row=r, column=1, value=opt)
            ws.cell(row=r, column=2, value=cnt)
            ws.cell(row=r, column=3,
                    value=f"=B{r}/SUM($B${data_start}:$B${data_start + len(q.options) - 1})")
            ws.cell(row=r, column=3).number_format = "0.0%"
            for c in (1, 2, 3):
                cell = ws.cell(row=r, column=c)
                cell.font = Font(name="Calibri", size=11)
                cell.border = Border(bottom=THIN)
                if i % 2:
                    cell.fill = PatternFill("solid", fgColor="F8FAFC")

        data_end = data_start + len(q.options) - 1

        # Total row
        total_row = data_end + 1
        ws.cell(row=total_row, column=1, value="Total").font = Font(
            bold=True, name="Calibri", size=11)
        ws.cell(row=total_row, column=2,
                value=f"=SUM(B{data_start}:B{data_end})").font = Font(
            bold=True, name="Calibri", size=11)
        for c in (1, 2, 3):
            ws.cell(row=total_row, column=c).border = Border(top=ACCENT_BOTTOM)

        # Native Excel chart
        chart_type = q.chart_type
        if chart_type in ("pie",):
            chart = PieChart()
        elif chart_type in ("donut", "doughnut"):
            chart = DoughnutChart()
        elif chart_type == "line":
            chart = LineChart()
        else:
            chart = BarChart()
            chart.type = "col"
            chart.style = 11

        chart.title = q.text
        chart.height = 10
        chart.width = 20
        cats = Reference(ws, min_col=1, min_row=data_start,
                         max_row=data_end)
        values = Reference(ws, min_col=2, min_row=data_start - 1,
                           max_row=data_end)   # incl header for legend
        chart.add_data(values, titles_from_data=True)
        chart.set_categories(cats)
        chart.dataLabels = DataLabelList(showVal=True)
        ws.add_chart(chart, f"E{row}")

    # ── Scale: stats + histogram bins + chart ─────────────────────────────
    elif q.qtype == "scale" and q.scale_values:
        ws.cell(row=row, column=1, value="Statistic")
        ws.cell(row=row, column=2, value="Value")
        _style_header_row(ws, row, 2)
        stats = [
            ("Average", q.average or 0),
            ("Median",  sorted(q.scale_values)[len(q.scale_values) // 2]),
            ("Min",     min(q.scale_values)),
            ("Max",     max(q.scale_values)),
            ("Count",   len(q.scale_values)),
        ]
        for i, (label, val) in enumerate(stats):
            r = row + 1 + i
            ws.cell(row=r, column=1, value=label).font = Font(
                bold=True, name="Calibri", size=11)
            ws.cell(row=r, column=2, value=val).font = Font(
                name="Calibri", size=11)
            if isinstance(val, float):
                ws.cell(row=r, column=2).number_format = "0.00"

        # Histogram
        hist_row = row + len(stats) + 3
        ws.cell(row=hist_row, column=1, value="Score")
        ws.cell(row=hist_row, column=2, value="Count")
        _style_header_row(ws, hist_row, 2)
        bin_counts = Counter(int(round(v)) for v in q.scale_values)
        for i in range(1, 11):
            r = hist_row + i
            ws.cell(row=r, column=1, value=i)
            ws.cell(row=r, column=2, value=bin_counts.get(i, 0))

        chart = BarChart()
        chart.type = "col"
        chart.title = q.text
        chart.height = 10
        chart.width = 20
        cats = Reference(ws, min_col=1, min_row=hist_row + 1,
                         max_row=hist_row + 10)
        values = Reference(ws, min_col=2, min_row=hist_row,
                           max_row=hist_row + 10)
        chart.add_data(values, titles_from_data=True)
        chart.set_categories(cats)
        chart.dataLabels = DataLabelList(showVal=True)
        ws.add_chart(chart, f"E{row}")

    # ── Word cloud: frequency table ───────────────────────────────────────
    elif q.qtype == "word" and q.words:
        ws.cell(row=row, column=1, value="Word")
        ws.cell(row=row, column=2, value="Mentions")
        _style_header_row(ws, row, 2)
        freqs = Counter(w.lower().strip() for w in q.words if w.strip())
        sorted_items = freqs.most_common()
        for i, (w, c) in enumerate(sorted_items):
            r = row + 1 + i
            ws.cell(row=r, column=1, value=w).font = Font(
                name="Calibri", size=11)
            ws.cell(row=r, column=2, value=c).font = Font(
                name="Calibri", size=11)
            if i % 2:
                for cc in (1, 2):
                    ws.cell(row=r, column=cc).fill = PatternFill(
                        "solid", fgColor="F8FAFC")

        # bar chart for top 15
        top_n = min(15, len(sorted_items))
        chart = BarChart()
        chart.type = "bar"   # horizontal
        chart.title = f"Top {top_n} words"
        chart.height = 12
        chart.width = 20
        cats = Reference(ws, min_col=1, min_row=row + 1,
                         max_row=row + top_n)
        values = Reference(ws, min_col=2, min_row=row,
                           max_row=row + top_n)
        chart.add_data(values, titles_from_data=True)
        chart.set_categories(cats)
        ws.add_chart(chart, f"E{row}")

    # ── Open text: verbatim list ─────────────────────────────────────────
    elif q.qtype == "open" and q.open_answers:
        ws.cell(row=row, column=1, value="#")
        ws.cell(row=row, column=2, value="Response")
        _style_header_row(ws, row, 2)
        for i, ans in enumerate(q.open_answers, 1):
            r = row + i
            ws.cell(row=r, column=1, value=i)
            ws.cell(row=r, column=2, value=ans)
            ws.cell(row=r, column=2).alignment = Alignment(
                wrap_text=True, vertical="top")
            for c in (1, 2):
                ws.cell(row=r, column=c).font = Font(
                    name="Calibri", size=11)
                if i % 2 == 0:
                    ws.cell(row=r, column=c).fill = PatternFill(
                        "solid", fgColor="F8FAFC")

    # Column widths
    for col_letter, w in zip("ABCDEFGHIJ", [44, 12, 12, 4, 30, 30, 30, 4, 4, 4]):
        ws.column_dimensions[col_letter].width = w


def build_excel_report(data: ReportData) -> bytes:
    wb = Workbook()
    _add_summary_sheet(wb, data)
    for i, q in enumerate(data.questions, 1):
        _add_question_sheet(wb, i, q)

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()
