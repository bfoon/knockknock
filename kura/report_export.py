"""Word, Excel, PDF and standalone HTML exports for Kura cleaned datasets."""

from __future__ import annotations

import html
import io
import json
from datetime import datetime

import pandas as pd
from django.http import HttpResponse

from .analytics import dashboard_payload, run_dataframe


def _filename(run, ext):
    code = run.pipeline.survey.code
    return f"{code}_cleaning_run_{run.id}.{ext}"


# ── repeat groups ────────────────────────────────────────────────────
# A repeat group (household members, plots, livestock…) lives in ONE cell
# of the parent row as a list of dicts. Excel can't hold that, so the
# exports below split each repeat group onto its own sheet, carrying the
# parent's identifiers so the two can be joined back together.

# Identifier columns copied onto every child row, in preference order.
PARENT_KEYS = ["_row_number", "_submission_id", "_uuid"]

# How many parent context columns (region, district…) to carry across
# so the child sheet is useful on its own without a manual join.
MAX_PARENT_CONTEXT = 6


def _sheet_name(base, used=None, limit=31):
    """Excel sheet names: <=31 chars, no []:*?/\\ , and must be unique."""
    clean = "".join("_" if ch in "[]:*?/\\" else ch for ch in str(base)).strip()
    clean = (clean or "Sheet")[:limit]
    if used is None:
        return clean
    name, n = clean, 2
    while name.lower() in used:
        suffix = f"_{n}"
        name = clean[: limit - len(suffix)] + suffix
        n += 1
    used.add(name.lower())
    return name


def _is_repeat_cell(value):
    return isinstance(value, list) and (not value or isinstance(value[0], dict))


def repeat_columns(df):
    """Columns that hold repeat-group answers, in dataframe order."""
    found = []
    for c in df.columns:
        sample = df[c].dropna().head(50)
        if sample.empty:
            continue
        hits = sum(1 for v in sample if _is_repeat_cell(v))
        if hits >= max(1, int(len(sample) * 0.6)):
            found.append(c)
    return found


def _parent_key_columns(df):
    return [c for c in PARENT_KEYS if c in df.columns]


def _parent_context_columns(df, repeats, keys):
    """Plain scalar parent columns worth carrying onto the child sheet.

    Ordered so that good grouping columns (region, district — few distinct
    values) come first, because the pivot uses the first one as its rows.
    A near-unique column like a household code groups nothing useful.
    """
    skip = set(repeats) | set(keys)
    candidates = []
    for c in df.columns:
        if c in skip or c.startswith("_"):
            continue
        if df[c].map(lambda v: isinstance(v, (list, dict))).any():
            continue
        candidates.append(c)

    total = max(1, len(df))

    def rank(col):
        distinct = df[col].astype(str).nunique()
        # 2..~15 distinct values makes a good grouping column
        good = 2 <= distinct <= max(15, total * 0.5)
        return (0 if good else 1, distinct)

    candidates.sort(key=rank)
    return candidates[:MAX_PARENT_CONTEXT]


def flatten_repeat(df, column, keys=None, context=None, with_fields=False):
    """One row per repeat ITEM, carrying the parent's identifiers.

    A parent with an empty repeat contributes no rows — the parent sheet
    still holds it, so nothing is lost.

    with_fields=True returns (frame, child_field_names) so callers can tell
    the group's own fields apart from the carried parent columns.
    """
    keys = keys if keys is not None else _parent_key_columns(df)
    context = context or []
    carry = [c for c in list(keys) + list(context) if c in df.columns]

    records, child_fields = [], []
    for _, parent in df.iterrows():
        items = parent[column]
        if not isinstance(items, list):
            continue
        for position, item in enumerate(items, start=1):
            if not isinstance(item, dict):
                item = {"value": item}
            record = {c: parent[c] for c in carry}
            record[f"{column}_index"] = position
            for k, v in item.items():
                # never let a child field clobber a parent identifier
                key = f"{column}_{k}" if k in record else k
                record[key] = json.dumps(v, ensure_ascii=False) if isinstance(v, (list, dict)) else v
                if key not in child_fields:
                    child_fields.append(key)
            records.append(record)

    if not records:
        return (None, []) if with_fields else None

    frame = pd.DataFrame(records)
    ordered = [c for c in carry + [f"{column}_index"] if c in frame.columns]
    rest = [c for c in frame.columns if c not in ordered]
    frame = frame[ordered + rest]
    return (frame, child_fields) if with_fields else frame


def _pivot_candidates(frame, exclude):
    """Split a child frame into (categorical, numeric) columns worth pivoting.

    Categorical candidates are ordered by how well they group: a field like
    sex or crop (few repeated values) is useful as a cross-tab column, while
    a near-unique field like a person's name would produce one column per
    row and tell you nothing. Anything with almost as many distinct values
    as rows is rejected outright.
    """
    categorical, numeric = [], []
    total = max(1, len(frame))
    for c in frame.columns:
        if c in exclude or c.startswith("_"):
            continue
        values = frame[c].dropna()
        if values.empty:
            continue
        as_num = pd.to_numeric(values, errors="coerce")
        if as_num.notna().sum() >= len(values) * 0.8:
            numeric.append(c)
            continue
        distinct = values.astype(str).nunique()
        # Must actually group: reject a field that is near-unique per row
        # (a person's name would give one column per row). The ratio test
        # only bites once there are enough rows for it to mean something.
        if distinct < 2 or distinct > 25:
            continue
        if total >= 8 and distinct > total * 0.6:
            continue
        if total < 8 and distinct >= total:
            continue
        categorical.append((distinct, c))

    categorical.sort()          # fewest distinct values first
    return [c for _, c in categorical], numeric


def repeat_pivot(frame, column, parent_context=None, child_fields=None):
    """A ready-made cross-tab of a repeat group.

    Rows = the first parent context column (region, district…) when we have
    one, else the repeat index. Columns = the child's own first categorical
    field (sex, relationship…). Values = count of items, plus the mean of
    the child's first numeric field when there is one.

    Only the repeat group's OWN fields are used as the measure and the
    column split — the parent columns are carried for joining, and
    cross-tabbing them against each other would be meaningless.

    Returns None when the data can't support a meaningful cross-tab.
    """
    if frame is None or frame.empty:
        return None

    parent_context = [c for c in (parent_context or []) if c in frame.columns]
    index_col = parent_context[0] if parent_context else f"{column}_index"
    if index_col not in frame.columns:
        return None

    # candidates are the child's own fields only
    carried = set(PARENT_KEYS) | set(parent_context) | {f"{column}_index"}
    if child_fields is not None:
        pool = [c for c in child_fields if c in frame.columns]
    else:
        pool = [c for c in frame.columns if c not in carried and not c.startswith("_")]
    if not pool:
        return None

    categorical, numeric = _pivot_candidates(frame[pool], set())

    try:
        if categorical:
            column_field = categorical[0]
            if numeric:
                measure = numeric[0]
                table = pd.pivot_table(
                    frame, index=index_col, columns=column_field,
                    values=measure, aggfunc=["count", "mean"], fill_value=0,
                )
                table.columns = [f"{a} of {measure} — {b}" for a, b in table.columns]
            else:
                table = pd.pivot_table(
                    frame, index=index_col, columns=column_field,
                    values=f"{column}_index", aggfunc="count", fill_value=0,
                )
                table.columns = [f"count — {c}" for c in table.columns]
        elif numeric:
            measure = numeric[0]
            table = frame.groupby(index_col)[measure].agg(["count", "mean", "min", "max"])
            table.columns = [f"{a} of {measure}" for a in table.columns]
        else:
            table = frame.groupby(index_col).size().to_frame("count")
    except (KeyError, ValueError, TypeError):
        return None

    if table.empty:
        return None
    table = table.round(2)
    table.insert(0, index_col, table.index)
    return table.reset_index(drop=True)


def _excel_safe(df):
    """Serialize anything Excel can't store, so no value is silently lost."""
    out = df.copy()
    for c in out.columns:
        out[c] = out[c].apply(
            lambda v: json.dumps(v, ensure_ascii=False)
            if isinstance(v, (list, dict)) else v
        )
    return out


def _write_repeat_sheets(writer, df, used_names):
    """Write one sheet per repeat group, plus a pivot sheet for each.

    Returns the list of repeat column names handled.
    """
    repeats = repeat_columns(df)
    if not repeats:
        return []

    keys = _parent_key_columns(df)
    context = _parent_context_columns(df, repeats, keys)

    for column in repeats:
        child, child_fields = flatten_repeat(
            df, column, keys=keys, context=context, with_fields=True)
        if child is None or child.empty:
            continue
        child.to_excel(
            writer, sheet_name=_sheet_name(column, used_names), index=False)

        pivot = repeat_pivot(child, column, parent_context=context,
                             child_fields=child_fields)
        if pivot is not None:
            pivot.to_excel(
                writer,
                sheet_name=_sheet_name(f"{column} pivot", used_names),
                index=False,
            )
    return repeats


def export_clean_csv(run):
    """Just the cleaned dataset as CSV — no report, no audit sheets."""
    import csv

    df = run_dataframe(run)
    out = io.StringIO()
    writer = csv.writer(out)
    columns = list(df.columns)
    writer.writerow(columns)
    for _, row in df.iterrows():
        values = []
        for c in columns:
            v = row[c]
            if isinstance(v, (list, dict)):
                values.append(json.dumps(v, ensure_ascii=False))
            elif v is None or (isinstance(v, float) and v != v):  # NaN check
                values.append("")
            else:
                values.append(v)
        writer.writerow(values)
    response = HttpResponse(
        out.getvalue().encode("utf-8-sig"),
        content_type="text/csv; charset=utf-8",
    )
    response["Content-Disposition"] = (
        f'attachment; filename="{run.pipeline.survey.code}_clean_data_run_{run.id}.csv"'
    )
    return response


def export_clean_excel(run):
    """The cleaned dataset as an Excel workbook.

    Sheet 1 is the parent rows. Every repeat group (household members,
    plots…) gets its own sheet keyed back to the parent, followed by a
    ready-made pivot of that group.
    """
    df = run_dataframe(run)
    if "_row_number" not in df.columns:
        df.insert(0, "_row_number", range(1, len(df) + 1))

    out = io.BytesIO()
    used_names = set()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        _excel_safe(df).to_excel(
            writer, sheet_name=_sheet_name("Clean Data", used_names), index=False)
        _write_repeat_sheets(writer, df, used_names)

    response = HttpResponse(
        out.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = (
        f'attachment; filename="{run.pipeline.survey.code}_clean_data_run_{run.id}.xlsx"'
    )
    return response


def export_excel(run):
    df = run_dataframe(run)
    if "_row_number" not in df.columns:
        df.insert(0, "_row_number", range(1, len(df) + 1))

    changes = list(run.changes.values(
        "row_number", "field", "change_type", "old_value", "new_value", "detail"
    ))
    steps = list(run.pipeline.steps.values(
        "order", "name", "operation", "enabled", "config", "note"
    ))
    stats = dashboard_payload(run)

    out = io.BytesIO()
    used_names = set()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        _excel_safe(df).to_excel(
            writer, sheet_name=_sheet_name("Cleaned Data", used_names), index=False)

        # one sheet per repeat group + its pivot, right after the parent data
        _write_repeat_sheets(writer, df, used_names)

        pd.DataFrame(changes).to_excel(
            writer, sheet_name=_sheet_name("Change Log", used_names), index=False)
        pd.DataFrame(steps).to_excel(
            writer, sheet_name=_sheet_name("Pipeline Steps", used_names), index=False)
        pd.DataFrame(stats.get("numeric", {})).T.to_excel(
            writer, sheet_name=_sheet_name("Summary Statistics", used_names))
        corr = stats.get("correlation", {})
        if corr.get("columns"):
            pd.DataFrame(
                corr["matrix"], index=corr["columns"], columns=corr["columns"]
            ).to_excel(writer, sheet_name=_sheet_name("Correlation Matrix", used_names))
    response = HttpResponse(
        out.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{_filename(run, "xlsx")}"'
    return response


def _report_context(run):
    return {
        "survey": run.pipeline.survey,
        "run": run,
        "stats": dashboard_payload(run),
        "rows": list(run.records.order_by("row_number")[:100]),
        "changes": list(run.changes.select_related("step").order_by("id")[:500]),
    }


def export_html(run):
    ctx = _report_context(run)
    stats = ctx["stats"]
    numeric_rows = "".join(
        f"<tr><td>{html.escape(name)}</td><td>{v.get('count')}</td>"
        f"<td>{v.get('mean')}</td><td>{v.get('median')}</td>"
        f"<td>{v.get('min')}</td><td>{v.get('max')}</td></tr>"
        for name, v in stats.get("numeric", {}).items()
    )
    doc = f"""<!doctype html>
<html><head><meta charset="utf-8"><title>{html.escape(ctx['survey'].title)} report</title>
<style>
body{{font:15px Arial,sans-serif;color:#202124;max-width:1100px;margin:40px auto;padding:0 25px}}
h1,h2{{color:#30205f}} .meta{{background:#f3f0ff;padding:16px;border-radius:10px}}
table{{border-collapse:collapse;width:100%;margin:15px 0}}th,td{{border:1px solid #ddd;padding:8px;text-align:left}}
th{{background:#30205f;color:white}} .small{{font-size:12px;color:#666}}
</style></head><body>
<h1>{html.escape(ctx['survey'].title)}</h1>
<div class="meta"><b>Cleaning report</b><br>
Pipeline: {html.escape(run.pipeline.name)}<br>
Run: {run.id} · Rows: {run.result_count} · Changes: {run.changes.count()}<br>
Generated: {datetime.utcnow().isoformat()} UTC</div>
<h2>Executive summary</h2>
<p>The cleaned result contains {run.result_count} rows and {run.column_count} columns.
The pipeline removed {run.excluded_count} rows and recorded {run.changes.count()} auditable changes.</p>
<h2>Pipeline steps</h2>
<ol>{''.join(f"<li><b>{html.escape(s.name)}</b> — {html.escape(s.operation)}</li>" for s in run.pipeline.steps.order_by("order"))}</ol>
<h2>Numeric summary</h2>
<table><thead><tr><th>Variable</th><th>N</th><th>Mean</th><th>Median</th><th>Min</th><th>Max</th></tr></thead>
<tbody>{numeric_rows}</tbody></table>
<h2>Correlation matrix</h2>
<pre>{html.escape(json.dumps(stats.get("correlation", {}), indent=2))}</pre>
<p class="small">Generated by Kura Data Studio.</p>
</body></html>"""
    response = HttpResponse(doc, content_type="text/html; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{_filename(run, "html")}"'
    return response


def export_word(run):
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word export.") from exc

    ctx = _report_context(run)
    stats = ctx["stats"]
    doc = Document()
    doc.add_heading(ctx["survey"].title, 0)
    doc.add_paragraph(
        f"Cleaning report · Pipeline: {run.pipeline.name} · Run #{run.id}"
    )
    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        f"The cleaned result contains {run.result_count} rows and "
        f"{run.column_count} columns. {run.excluded_count} rows were removed "
        f"and {run.changes.count()} changes were recorded."
    )
    doc.add_heading("Pipeline steps", level=1)
    for step in run.pipeline.steps.order_by("order"):
        doc.add_paragraph(
            f"{step.order}. {step.name} ({step.operation})",
            style="List Number",
        )
    doc.add_heading("Numeric summary", level=1)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    for i, label in enumerate(["Variable", "N", "Mean", "Median", "Min", "Max"]):
        hdr[i].text = label
    for name, values in stats.get("numeric", {}).items():
        cells = table.add_row().cells
        data = [
            name, values.get("count"), values.get("mean"),
            values.get("median"), values.get("min"), values.get("max"),
        ]
        for i, value in enumerate(data):
            cells[i].text = "" if value is None else str(value)
    out = io.BytesIO()
    doc.save(out)
    response = HttpResponse(
        out.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{_filename(run, "docx")}"'
    return response


def export_pdf(run):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError("reportlab is required for PDF export.") from exc

    ctx = _report_context(run)
    stats = ctx["stats"]
    out = io.BytesIO()
    doc = SimpleDocTemplate(
        out, pagesize=A4, rightMargin=16*mm, leftMargin=16*mm,
        topMargin=16*mm, bottomMargin=16*mm,
    )
    styles = getSampleStyleSheet()
    story = [
        Paragraph(ctx["survey"].title, styles["Title"]),
        Paragraph(f"Cleaning report · Pipeline {run.pipeline.name} · Run #{run.id}", styles["Normal"]),
        Spacer(1, 10),
        Paragraph("Executive summary", styles["Heading1"]),
        Paragraph(
            f"The cleaned result contains {run.result_count} rows and {run.column_count} columns. "
            f"{run.excluded_count} rows were removed and {run.changes.count()} changes were recorded.",
            styles["BodyText"],
        ),
        Paragraph("Numeric summary", styles["Heading1"]),
    ]
    table_data = [["Variable", "N", "Mean", "Median", "Min", "Max"]]
    for name, v in stats.get("numeric", {}).items():
        table_data.append([
            name, v.get("count"), v.get("mean"), v.get("median"), v.get("min"), v.get("max")
        ])
    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#30205f")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), .5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(table)
    doc.build(story)
    response = HttpResponse(out.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{_filename(run, "pdf")}"'
    return response