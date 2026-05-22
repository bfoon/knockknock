"""
Excel / Word export of the post-event feedback report.

The organizer's results page (feedback_views.feedback_results) renders
the same numbers on screen; this module produces downloadable .xlsx and
.docx versions of that report.

Layout of each export mirrors the results page:

  * A header block (event title, response count, generated-at).
  * A "Question summaries" section — one block per question with the
    same aggregates the HTML shows (rating averages + star counts,
    multiple-choice option counts, yes/no counts, open-text answers).
  * An "Individual responses" section — one block per submission with
    every answer.

The aggregation itself is factored into `build_report_context()` so
the two builders stay thin and there is a single source of truth for
the numbers. It is intentionally close to the logic in
`feedback_results` and can be reused there later if desired.

Both builders return raw bytes; the view layer wraps them in an
HttpResponse with the right content-type and Content-Disposition.
"""

import io
from collections import Counter
from datetime import datetime

from django.utils import timezone

from .feedback_models import FeedbackQuestion, FeedbackAnswer


# ─────────────────────────── Aggregation ───────────────────────────

def build_report_context(event, survey):
    """
    Compute everything the exports need from a survey.

    Returns a dict:
      {
        "event": event,
        "survey": survey,
        "generated_at": datetime (aware),
        "response_count": int,
        "question_summaries": [ {...}, ... ],   # same shape as the view
        "responses": [ {                        # per-submission rows
            "name": str,
            "submitted_at": datetime,
            "answers": [ {"question": str, "value": str}, ... ],
          }, ... ],
      }

    `question_summaries` entries carry a `kind` of one of:
    separator / open_text / rate_1_5 / multiple_choice / yes_no — the
    same vocabulary the results template already branches on.
    """
    responses_qs = (
        survey.responses
        .prefetch_related("answers__question")
        .order_by("-submitted_at")
    )

    question_summaries = []
    for q in survey.questions.all().order_by("order", "id"):
        if q.is_separator():
            question_summaries.append({"question": q, "kind": "separator"})
            continue

        answers = FeedbackAnswer.objects.filter(question=q)

        if q.question_type == FeedbackQuestion.TYPE_OPEN_TEXT:
            entries = [
                {"name": a.response.display_name(), "text": a.text_answer}
                for a in answers.select_related("response")
                if a.text_answer
            ]
            question_summaries.append({
                "question": q,
                "kind": "open_text",
                "entries": entries,
                "count": len(entries),
            })

        elif q.question_type == FeedbackQuestion.TYPE_RATE_1_5:
            buckets = Counter()
            ratings = []
            for a in answers:
                if a.rating:
                    buckets[a.rating] += 1
                    ratings.append(a.rating)
            total = sum(buckets.values())
            avg = (sum(ratings) / len(ratings)) if ratings else 0
            rows = [(star, buckets.get(star, 0)) for star in range(1, 6)]
            question_summaries.append({
                "question": q,
                "kind": "rate_1_5",
                "rows": rows,
                "count": total,
                "average": round(avg, 2) if avg else 0,
            })

        elif q.question_type == FeedbackQuestion.TYPE_MULTIPLE_CHOICE:
            buckets = Counter()
            for a in answers:
                if a.choice_answer:
                    buckets[a.choice_answer] += 1
            defined = q.cleaned_choices()
            rows = [(opt, buckets.get(opt, 0)) for opt in defined]
            # Any stored answers not in the current option list.
            for k, v in buckets.items():
                if k not in defined:
                    rows.append((k, v))
            question_summaries.append({
                "question": q,
                "kind": "multiple_choice",
                "rows": rows,
                "count": sum(v for _, v in rows),
            })

        elif q.question_type == FeedbackQuestion.TYPE_YES_NO:
            yes = answers.filter(bool_answer=True).count()
            no = answers.filter(bool_answer=False).count()
            question_summaries.append({
                "question": q,
                "kind": "yes_no",
                "yes": yes,
                "no": no,
                "count": yes + no,
            })

    responses = []
    for r in responses_qs:
        responses.append({
            "name": r.display_name(),
            "submitted_at": r.submitted_at,
            "answers": [
                {
                    "question": a.question.text,
                    "value": a.display_value() or "—",
                }
                for a in r.answers.all()
            ],
        })

    return {
        "event": event,
        "survey": survey,
        "generated_at": timezone.now(),
        "response_count": responses_qs.count(),
        "question_summaries": question_summaries,
        "responses": responses,
    }


def export_filename(event, ext):
    """A safe-ish download filename like `feedback-my-event.xlsx`."""
    slug = "".join(
        c if c.isalnum() else "-"
        for c in (event.title or "event").lower()
    ).strip("-") or "event"
    # Collapse runs of dashes.
    while "--" in slug:
        slug = slug.replace("--", "-")
    return f"feedback-{slug}.{ext}"


def _pct(part, whole):
    return round((part / whole) * 100) if whole else 0


# ─────────────────────────── Excel builder ───────────────────────────

def build_feedback_xlsx(event, survey):
    """
    Return .xlsx bytes for the feedback report.

    Three sheets:
      Summary    — header info + one section per question with counts.
      Responses  — one row per (response, question) — tidy/long format,
                   easy to pivot.
      Open text  — every free-text answer with its author.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    ctx = build_report_context(event, survey)

    wb = Workbook()

    # Shared styles.
    title_font = Font(size=15, bold=True, color="1F2937")
    h2_font = Font(size=12, bold=True, color="FFFFFF")
    h2_fill = PatternFill("solid", fgColor="4F46E5")
    q_font = Font(size=11, bold=True, color="111827")
    head_font = Font(bold=True, color="FFFFFF")
    head_fill = PatternFill("solid", fgColor="6B7280")
    dim_font = Font(size=10, color="6B7280")
    thin = Side(style="thin", color="D1D5DB")
    box = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── Sheet 1: Summary ────────────────────────────────────────
    ws = wb.active
    ws.title = "Summary"
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 40

    row = 1
    ws.cell(row=row, column=1, value="Feedback report").font = title_font
    row += 1
    ws.cell(row=row, column=1, value=event.title).font = Font(size=12, bold=True)
    row += 1
    ws.cell(row=row, column=1,
            value=f"{ctx['response_count']} response(s)").font = dim_font
    row += 1
    ws.cell(
        row=row, column=1,
        value=f"Generated {ctx['generated_at'].strftime('%Y-%m-%d %H:%M')}",
    ).font = dim_font
    row += 2

    def section_header(label):
        nonlocal row
        c = ws.cell(row=row, column=1, value=label)
        c.font = h2_font
        c.fill = h2_fill
        for col in range(2, 5):
            ws.cell(row=row, column=col).fill = h2_fill
        row += 1

    def table_header(cols):
        nonlocal row
        for i, name in enumerate(cols, start=1):
            c = ws.cell(row=row, column=i, value=name)
            c.font = head_font
            c.fill = head_fill
            c.border = box
        row += 1

    section_header("Question summaries")
    row += 1

    if not ctx["question_summaries"]:
        ws.cell(row=row, column=1, value="No questions on this survey.").font = dim_font
        row += 1

    for s in ctx["question_summaries"]:
        kind = s["kind"]
        q = s["question"]

        if kind == "separator":
            c = ws.cell(row=row, column=1, value=f"— {q.text} —")
            c.font = Font(size=11, bold=True, italic=True, color="4F46E5")
            row += 2
            continue

        # Question label line.
        c = ws.cell(row=row, column=1, value=q.text)
        c.font = q_font
        c.alignment = Alignment(wrap_text=True, vertical="top")
        meta = q.get_question_type_display()
        if kind == "rate_1_5":
            meta += f"  ·  avg {s['average']}/5"
        ws.cell(row=row, column=2,
                value=f"{s['count']} response(s)").font = dim_font
        ws.cell(row=row, column=3, value=meta).font = dim_font
        row += 1

        if kind == "open_text":
            if s["entries"]:
                table_header(["Respondent", "Answer"])
                for e in s["entries"]:
                    ws.cell(row=row, column=1, value=e["name"]).border = box
                    ac = ws.cell(row=row, column=2, value=e["text"])
                    ac.border = box
                    ac.alignment = Alignment(wrap_text=True, vertical="top")
                    # Let the wide answer span the remaining columns visually.
                    row += 1
            else:
                ws.cell(row=row, column=1,
                        value="No answers submitted.").font = dim_font
                row += 1

        elif kind in ("rate_1_5", "multiple_choice"):
            table_header(["Option", "Count", "Percent", ""])
            total = s["count"]
            for label, count in s["rows"]:
                disp = f"{label} ★" if kind == "rate_1_5" else str(label)
                ws.cell(row=row, column=1, value=disp).border = box
                ws.cell(row=row, column=2, value=count).border = box
                pc = ws.cell(row=row, column=3, value=_pct(count, total) / 100)
                pc.number_format = "0%"
                pc.border = box
                # Tiny text bar for quick visual scan.
                bar = "█" * int(round(_pct(count, total) / 5))
                ws.cell(row=row, column=4, value=bar).font = Font(color="6366F1")
                row += 1

        elif kind == "yes_no":
            table_header(["Option", "Count", "Percent", ""])
            total = s["count"]
            for label, count in (("Yes", s["yes"]), ("No", s["no"])):
                ws.cell(row=row, column=1, value=label).border = box
                ws.cell(row=row, column=2, value=count).border = box
                pc = ws.cell(row=row, column=3, value=_pct(count, total) / 100)
                pc.number_format = "0%"
                pc.border = box
                bar = "█" * int(round(_pct(count, total) / 5))
                ws.cell(row=row, column=4, value=bar).font = Font(color="6366F1")
                row += 1

        row += 1  # gap between questions

    # ── Sheet 2: Responses (long format) ────────────────────────
    ws2 = wb.create_sheet("Responses")
    ws2.sheet_view.showGridLines = False
    headers = ["#", "Respondent", "Submitted", "Question", "Answer"]
    widths = [6, 24, 20, 40, 50]
    for i, (name, w) in enumerate(zip(headers, widths), start=1):
        c = ws2.cell(row=1, column=i, value=name)
        c.font = head_font
        c.fill = head_fill
        c.border = box
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.freeze_panes = "A2"

    r2 = 2
    for idx, resp in enumerate(ctx["responses"], start=1):
        submitted = resp["submitted_at"]
        submitted_str = (
            timezone.localtime(submitted).strftime("%Y-%m-%d %H:%M")
            if submitted else ""
        )
        if not resp["answers"]:
            ws2.cell(row=r2, column=1, value=idx)
            ws2.cell(row=r2, column=2, value=resp["name"])
            ws2.cell(row=r2, column=3, value=submitted_str)
            ws2.cell(row=r2, column=4, value="(no answers)").font = dim_font
            r2 += 1
            continue
        for a in resp["answers"]:
            ws2.cell(row=r2, column=1, value=idx)
            ws2.cell(row=r2, column=2, value=resp["name"])
            ws2.cell(row=r2, column=3, value=submitted_str)
            ws2.cell(row=r2, column=4, value=a["question"]).alignment = \
                Alignment(wrap_text=True, vertical="top")
            ws2.cell(row=r2, column=5, value=a["value"]).alignment = \
                Alignment(wrap_text=True, vertical="top")
            r2 += 1

    if r2 == 2:
        ws2.cell(row=2, column=1, value="No responses yet.").font = dim_font

    # ── Sheet 3: Open text only ─────────────────────────────────
    ws3 = wb.create_sheet("Open text")
    ws3.sheet_view.showGridLines = False
    for i, (name, w) in enumerate(
        zip(["Question", "Respondent", "Answer"], [34, 24, 60]), start=1
    ):
        c = ws3.cell(row=1, column=i, value=name)
        c.font = head_font
        c.fill = head_fill
        c.border = box
        ws3.column_dimensions[get_column_letter(i)].width = w
    ws3.freeze_panes = "A2"

    r3 = 2
    for s in ctx["question_summaries"]:
        if s["kind"] != "open_text":
            continue
        for e in s["entries"]:
            ws3.cell(row=r3, column=1, value=s["question"].text).alignment = \
                Alignment(wrap_text=True, vertical="top")
            ws3.cell(row=r3, column=2, value=e["name"])
            ws3.cell(row=r3, column=3, value=e["text"]).alignment = \
                Alignment(wrap_text=True, vertical="top")
            r3 += 1
    if r3 == 2:
        ws3.cell(row=2, column=1, value="No open-text answers.").font = dim_font

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─────────────────────────── Word builder ───────────────────────────

def build_feedback_docx(event, survey):
    """
    Return .docx bytes for the feedback report — a readable document
    version of the results page (summaries first, then every response).
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    ctx = build_report_context(event, survey)

    doc = Document()

    # Base style.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    INDIGO = RGBColor(0x4F, 0x46, 0xE5)
    GREY = RGBColor(0x6B, 0x72, 0x80)

    # ── Header ──────────────────────────────────────────────────
    title = doc.add_heading("Feedback report", level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(event.title)
    run.bold = True
    run.font.size = Pt(13)

    meta = doc.add_paragraph()
    mrun = meta.add_run(
        f"{ctx['response_count']} response(s)  ·  "
        f"Generated {ctx['generated_at'].strftime('%Y-%m-%d %H:%M')}"
    )
    mrun.font.color.rgb = GREY
    mrun.font.size = Pt(9)

    if survey.intro_text:
        ip = doc.add_paragraph(survey.intro_text)
        ip.italic = True

    # ── Question summaries ──────────────────────────────────────
    doc.add_heading("Question summaries", level=1)

    if not ctx["question_summaries"]:
        doc.add_paragraph("No questions on this survey.")

    def _bar_table(rows, total, star=False):
        """A 3-column count/percent/bar table."""
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Light Grid Accent 1"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tbl.rows[0].cells
        hdr[0].text = "Option"
        hdr[1].text = "Count"
        hdr[2].text = "Percent"
        for c in hdr:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
        for label, count in rows:
            cells = tbl.add_row().cells
            cells[0].text = (f"{label} ★" if star else str(label))
            cells[1].text = str(count)
            cells[2].text = f"{_pct(count, total)}%"
        return tbl

    for s in ctx["question_summaries"]:
        kind = s["kind"]
        q = s["question"]

        if kind == "separator":
            p = doc.add_paragraph()
            r = p.add_run(q.text.upper())
            r.bold = True
            r.font.color.rgb = INDIGO
            r.font.size = Pt(11)
            continue

        # Question heading.
        qh = doc.add_heading(level=2)
        qr = qh.add_run(q.text)

        info = doc.add_paragraph()
        meta_txt = f"{q.get_question_type_display()} · {s['count']} response(s)"
        if kind == "rate_1_5":
            meta_txt += f" · average {s['average']}/5"
        ir = info.add_run(meta_txt)
        ir.font.color.rgb = GREY
        ir.font.size = Pt(9)

        if kind == "open_text":
            if s["entries"]:
                for e in s["entries"]:
                    bullet = doc.add_paragraph(style="List Bullet")
                    bullet.add_run(e["text"])
                    who = bullet.add_run(f"  — {e['name']}")
                    who.italic = True
                    who.font.color.rgb = GREY
                    who.font.size = Pt(9)
            else:
                doc.add_paragraph("No answers submitted.")

        elif kind in ("rate_1_5", "multiple_choice"):
            _bar_table(s["rows"], s["count"], star=(kind == "rate_1_5"))

        elif kind == "yes_no":
            _bar_table(
                [("Yes", s["yes"]), ("No", s["no"])],
                s["count"],
            )

    # ── Individual responses ────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Individual responses", level=1)

    if not ctx["responses"]:
        doc.add_paragraph("No responses yet.")
    else:
        for idx, resp in enumerate(ctx["responses"], start=1):
            head = doc.add_heading(level=2)
            head.add_run(f"{idx}. {resp['name']}")
            submitted = resp["submitted_at"]
            if submitted:
                stamp = doc.add_paragraph()
                sr = stamp.add_run(
                    timezone.localtime(submitted).strftime("%a %d %b · %H:%M")
                )
                sr.font.color.rgb = GREY
                sr.font.size = Pt(9)

            if not resp["answers"]:
                doc.add_paragraph("(no answers)")
                continue

            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Light List Accent 1"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Question"
            hdr[1].text = "Answer"
            for c in hdr:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.bold = True
            for a in resp["answers"]:
                cells = tbl.add_row().cells
                cells[0].text = a["question"]
                cells[1].text = a["value"]
            # Column widths.
            for r in tbl.rows:
                r.cells[0].width = Inches(2.6)
                r.cells[1].width = Inches(3.6)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
