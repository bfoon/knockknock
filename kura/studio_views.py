"""HTTP endpoints for Kura's visual cleaning pipeline and analytics studio."""

from __future__ import annotations

import json

from django.contrib.auth.decorators import login_required
from django.db import transaction
from django.http import Http404, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, render
from django.views.decorators.http import require_GET, require_POST

from .analytics import dashboard_payload
from .models import (
    AnalysisDashboard,
    CleaningPipeline,
    CleaningRun,
    PipelineStep,
    Survey,
    UploadedDataset,
)
from .pipeline_engine import PipelineExecutor
from .report_export import (
    export_clean_csv,
    export_clean_excel,
    export_excel,
    export_html,
    export_pdf,
    export_word,
)

# Operations the pipeline engine supports that may not yet be listed in the
# model's OPERATION_CHOICES (no migration needed to use them). Merged with the
# model choices when validating and when listing operations for the studio UI.
EXTRA_OPERATIONS = [
    ("make_id", "Generate ID from columns"),
    ("concat_columns", "Combine columns into text"),
    ("extract_datetime", "Extract date/time parts"),
    # ── rows & shaping ──
    ("sort_rows", "Sort rows"),
    ("split_column", "Split column into parts"),
    ("replace_text", "Find and replace text"),
    ("bin_column", "Group numbers into bands"),
    ("clip_range", "Limit to a valid range"),
    ("round_numbers", "Round numbers"),
    ("group_aggregate", "Add group summary column"),
    ("rank_rows", "Rank rows"),
    ("running_total", "Running total"),
    ("flag_rows", "Mark rows (keep them)"),
    ("strip_accents", "Remove accents"),
    # ── machine learning ──
    ("cluster", "Find natural groups (ML)"),
    ("detect_anomalies", "Detect unusual rows (ML)"),
    ("predict_column", "Predict missing answers (ML)"),
    ("similar_duplicates", "Find near-duplicates (ML)"),
    ("reduce_dimensions", "Summarise many columns (ML)"),
]


def _all_operations():
    seen = {}
    for value, label in list(PipelineStep.OPERATION_CHOICES) + EXTRA_OPERATIONS:
        seen.setdefault(value, label)
    return list(seen.items())


def _own(request, code):
    survey = get_object_or_404(Survey, code=code.upper())
    if survey.owner_id != request.user.id:
        raise Http404
    return survey


def _json(request):
    try:
        return json.loads(request.body.decode("utf-8") or "{}")
    except (ValueError, UnicodeDecodeError):
        return None


def _survey_schema(survey):
    return (
        survey.current_version.schema
        if survey.current_version else survey.draft_schema
    ) or {}


def _scan_column_values(rows, cap=100, max_rows=2000):
    """Distinct values per column (capped), same shape the studio's friendly
    step editors expect. `rows` is a list of flat dicts."""
    column_values = {}
    for row in rows[:max_rows]:
        for key, value in (row or {}).items():
            if value in (None, "", []):
                continue
            bucket = column_values.setdefault(key, {})
            for v in (value if isinstance(value, list) else [value]):
                if isinstance(v, dict):
                    continue  # repeat-group items are browsed, not recoded
                sv = str(v)
                if len(sv) <= 120:
                    bucket[sv] = bucket.get(sv, 0) + 1
    return {
        col: [{"value": v, "count": n}
              for v, n in sorted(vals.items(), key=lambda kv: -kv[1])[:cap]]
        for col, vals in column_values.items()
    }


def _dataset_payload(ds, with_rows=False):
    payload = {
        "id": ds.id,
        "name": ds.name,
        "original_filename": ds.original_filename,
        "columns": ds.columns,
        "row_count": ds.row_count,
        "created_at": ds.created_at.isoformat(),
    }
    if with_rows:
        payload["rows"] = (ds.rows or [])[:500]
        payload["column_values"] = _scan_column_values(ds.rows or [])
    return payload


def _pipeline_payload(pipeline, with_steps=True):
    latest = pipeline.runs.filter(status="complete").order_by("-completed_at", "-id").first()
    payload = {
        "id": pipeline.id,
        "name": pipeline.name,
        "description": pipeline.description,
        "is_active": pipeline.is_active,
        "source": pipeline.source,
        "source_dataset_id": pipeline.source_dataset_id,
        "source_dataset_name": pipeline.source_dataset.name if pipeline.source_dataset_id else None,
        "step_count": pipeline.steps.count(),
        "updated_at": pipeline.updated_at.isoformat(),
        "latest_run": {
            "id": latest.id,
            "result_count": latest.result_count,
            "column_count": latest.column_count,
            "completed_at": latest.completed_at.isoformat() if latest.completed_at else None,
        } if latest else None,
    }
    if with_steps:
        payload["steps"] = [
            {
                "id": s.id,
                "order": s.order,
                "operation": s.operation,
                "name": s.name,
                "config": s.config,
                "enabled": s.enabled,
                "stop_on_error": s.stop_on_error,
                "note": s.note,
            }
            for s in pipeline.steps.order_by("order", "id")
        ]
    return payload


@login_required
def studio(request, code):
    survey = _own(request, code)
    pipeline = None
    # ?pipeline=<id> lets the data page (and bookmarks) open a specific
    # saved pipeline directly in the studio.
    requested = request.GET.get("pipeline")
    if requested:
        try:
            pipeline = survey.pipelines.filter(id=int(requested)).first()
        except (TypeError, ValueError):
            pipeline = None
    if pipeline is None:
        pipeline = survey.pipelines.order_by("-is_active", "-updated_at").first()
    if pipeline is None:
        pipeline = CleaningPipeline.objects.create(
            survey=survey,
            name="Main cleaning pipeline",
            created_by=request.user,
        )
    schema = _survey_schema(survey)
    columns = [
        q.get("name") for q in schema.get("questions", [])
        if q.get("name") and q.get("type") not in ("section", "note")
    ]
    return render(request, "kura/data_studio.html", {
        "survey": survey,
        "pipeline_json": json.dumps(_pipeline_payload(pipeline)),
        "columns_json": json.dumps(columns),
        "operations_json": json.dumps([
            {"value": value, "label": label}
            for value, label in _all_operations()
        ]),
    })


@login_required
@require_GET
def studio_bootstrap(request, code):
    survey = _own(request, code)
    pipelines = [
        _pipeline_payload(p)
        for p in survey.pipelines.prefetch_related("steps")
    ]
    runs = [
        {
            "id": r.id,
            "pipeline_id": r.pipeline_id,
            "pipeline_name": r.pipeline.name,
            "label": r.label,
            "status": r.status,
            "source_count": r.source_count,
            "result_count": r.result_count,
            "excluded_count": r.excluded_count,
            "column_count": r.column_count,
            "summary": r.summary,
            "error": r.error,
            "started_at": r.started_at.isoformat(),
            "completed_at": r.completed_at.isoformat() if r.completed_at else None,
        }
        for r in CleaningRun.objects.filter(pipeline__survey=survey).select_related("pipeline")[:30]
    ]
    raw_rows = [
        s.as_dict()
        for s in survey.submissions.select_related("form_version", "device")[:500]
    ]

    # For the friendly (no-JSON) step editors: the distinct values actually
    # present in each column, plus any labelled choices from the form schema.
    schema = _survey_schema(survey)
    schema_choices = {}
    for q in schema.get("questions", []):
        name = q.get("name")
        if name and q.get("choices"):
            schema_choices[name] = [
                {"value": str(c.get("value")), "label": c.get("label") or str(c.get("value"))}
                for c in q["choices"]
            ]

    merged_rows = [
        {**(row.get("answers") or {}), **(row.get("calculations") or {})}
        for row in raw_rows
    ]
    column_values = _scan_column_values(merged_rows)

    return JsonResponse({
        "ok": True,
        "pipelines": pipelines,
        "runs": runs,
        "raw_rows": raw_rows,
        "column_values": column_values,
        "schema_choices": schema_choices,
        "datasets": [
            _dataset_payload(d) for d in survey.datasets.all()
        ],
    })


@login_required
@require_GET
def pipeline_list(request, code):
    """Lightweight pipeline listing for the classic data page."""
    survey = _own(request, code)
    return JsonResponse({
        "ok": True,
        "pipelines": [
            _pipeline_payload(p, with_steps=False)
            for p in survey.pipelines.all()
        ],
    })


@login_required
@require_POST
def pipeline_save(request, code):
    survey = _own(request, code)
    data = _json(request)
    if data is None:
        return JsonResponse({"ok": False, "error": "JSON body required."}, status=400)

    pipeline_id = data.get("id")
    if pipeline_id:
        pipeline = get_object_or_404(CleaningPipeline, id=pipeline_id, survey=survey)
    else:
        pipeline = CleaningPipeline(survey=survey, created_by=request.user)

    pipeline.name = str(data.get("name") or "Cleaning pipeline")[:140]
    pipeline.description = str(data.get("description") or "")
    pipeline.is_active = bool(data.get("is_active", True))

    source = data.get("source") or "submissions"
    dataset_id = data.get("source_dataset_id")
    if source == "dataset" and dataset_id:
        dataset = survey.datasets.filter(id=dataset_id).first()
        if dataset is None:
            return JsonResponse({"ok": False, "error": "That uploaded file no longer exists."}, status=400)
        pipeline.source = "dataset"
        pipeline.source_dataset = dataset
    else:
        pipeline.source = "submissions"
        pipeline.source_dataset = None
    pipeline.save()

    steps = data.get("steps")
    if not isinstance(steps, list):
        return JsonResponse({"ok": False, "error": "steps must be a list."}, status=400)

    valid_ops = dict(_all_operations())
    with transaction.atomic():
        pipeline.steps.all().delete()
        objects = []
        for index, item in enumerate(steps):
            operation = item.get("operation")
            if operation not in valid_ops:
                return JsonResponse({
                    "ok": False,
                    "error": f"Unsupported operation: {operation}",
                }, status=400)
            objects.append(PipelineStep(
                pipeline=pipeline,
                order=index,
                operation=operation,
                name=str(item.get("name") or valid_ops[operation])[:140],
                config=item.get("config") if isinstance(item.get("config"), dict) else {},
                enabled=bool(item.get("enabled", True)),
                stop_on_error=bool(item.get("stop_on_error", True)),
                note=str(item.get("note") or ""),
            ))
        PipelineStep.objects.bulk_create(objects)

    pipeline.refresh_from_db()
    return JsonResponse({"ok": True, "pipeline": _pipeline_payload(pipeline)})


@login_required
@require_POST
def pipeline_duplicate(request, code, pipeline_id):
    """Clone a saved pipeline (steps included) so it can be adapted safely."""
    survey = _own(request, code)
    source = get_object_or_404(CleaningPipeline, id=pipeline_id, survey=survey)
    with transaction.atomic():
        clone = CleaningPipeline.objects.create(
            survey=survey,
            name=f"{source.name} (copy)"[:140],
            description=source.description,
            is_active=False,
            source=source.source,
            source_dataset=source.source_dataset,
            created_by=request.user,
        )
        PipelineStep.objects.bulk_create([
            PipelineStep(
                pipeline=clone,
                order=s.order,
                operation=s.operation,
                name=s.name,
                config=s.config,
                enabled=s.enabled,
                stop_on_error=s.stop_on_error,
                note=s.note,
            )
            for s in source.steps.order_by("order", "id")
        ])
    return JsonResponse({"ok": True, "pipeline": _pipeline_payload(clone)})


@login_required
@require_POST
def pipeline_delete(request, code, pipeline_id):
    """Delete a saved pipeline and its runs (raw submissions are untouched)."""
    survey = _own(request, code)
    pipeline = get_object_or_404(CleaningPipeline, id=pipeline_id, survey=survey)
    pipeline.delete()
    return JsonResponse({"ok": True})


# ─────────────────────────────────────────────────────────────────────
# Uploaded datasets (CSV / Excel as a pipeline source)
# ─────────────────────────────────────────────────────────────────────

MAX_DATASET_ROWS = 50_000
MAX_DATASET_COLUMNS = 300
MAX_DATASET_BYTES = 20 * 1024 * 1024  # 20 MB upload cap


@login_required
@require_POST
def dataset_upload(request, code):
    """Multipart upload of a CSV/Excel file → UploadedDataset.

    The whole file is parsed with pandas so CSV and Excel behave
    identically (headers, NA handling, mixed types)."""
    survey = _own(request, code)
    f = request.FILES.get("file")
    if f is None:
        return JsonResponse({"ok": False, "error": "Attach a CSV or Excel file."}, status=400)
    if f.size > MAX_DATASET_BYTES:
        return JsonResponse({"ok": False, "error": "File is larger than 20 MB."}, status=400)

    import pandas as pd

    name = (f.name or "upload").rsplit("/", 1)[-1]
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    try:
        if ext == "csv":
            df = pd.read_csv(f)
        elif ext in ("xlsx", "xls", "xlsm"):
            df = pd.read_excel(f)
        else:
            return JsonResponse({
                "ok": False,
                "error": "Only .csv, .xlsx, .xls and .xlsm files are supported.",
            }, status=400)
    except Exception as exc:
        return JsonResponse({"ok": False, "error": f"Could not read the file: {exc}"}, status=400)

    if df.empty:
        return JsonResponse({"ok": False, "error": "The file has no data rows."}, status=400)
    if len(df.columns) > MAX_DATASET_COLUMNS:
        return JsonResponse({"ok": False, "error": f"Too many columns (limit {MAX_DATASET_COLUMNS})."}, status=400)
    if len(df) > MAX_DATASET_ROWS:
        df = df.head(MAX_DATASET_ROWS)

    # Clean headers and convert to plain JSON rows (NaN → None).
    df.columns = [str(c).strip() or f"column_{i + 1}" for i, c in enumerate(df.columns)]
    df = df.where(pd.notna(df), None)
    rows = []
    for record in df.to_dict(orient="records"):
        rows.append({
            k: (v.isoformat() if hasattr(v, "isoformat") else v)
            for k, v in record.items()
        })

    ds = UploadedDataset.objects.create(
        survey=survey,
        name=str(request.POST.get("name") or name.rsplit(".", 1)[0])[:140],
        original_filename=name[:200],
        columns=list(df.columns),
        rows=rows,
        row_count=len(rows),
        uploaded_by=request.user,
    )
    return JsonResponse({"ok": True, "dataset": _dataset_payload(ds, with_rows=True)})


@login_required
@require_GET
def dataset_detail(request, code, dataset_id):
    """Columns, distinct values and a row preview for one uploaded file —
    fetched when the studio switches a pipeline onto that source."""
    survey = _own(request, code)
    ds = get_object_or_404(UploadedDataset, id=dataset_id, survey=survey)
    return JsonResponse({"ok": True, "dataset": _dataset_payload(ds, with_rows=True)})


@login_required
@require_POST
def dataset_delete(request, code, dataset_id):
    survey = _own(request, code)
    ds = get_object_or_404(UploadedDataset, id=dataset_id, survey=survey)
    # Pipelines pointing at it fall back to survey submissions.
    ds.pipelines.update(source="submissions", source_dataset=None)
    ds.delete()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def pipeline_run(request, code, pipeline_id):
    survey = _own(request, code)
    pipeline = get_object_or_404(CleaningPipeline, id=pipeline_id, survey=survey)
    data = _json(request) or {}
    run = CleaningRun.objects.create(
        pipeline=pipeline,
        label=str(data.get("label") or "")[:140],
        run_by=request.user,
    )
    try:
        PipelineExecutor(run).execute()
    except Exception as exc:
        return JsonResponse({
            "ok": False,
            "run_id": run.id,
            "error": str(exc),
        }, status=422)
    return JsonResponse({
        "ok": True,
        "run_id": run.id,
        "status": run.status,
        "summary": run.summary,
        "result_count": run.result_count,
        "column_count": run.column_count,
    })


@login_required
@require_GET
def run_rows(request, code, run_id):
    survey = _own(request, code)
    run = get_object_or_404(CleaningRun, id=run_id, pipeline__survey=survey)
    limit = min(2000, max(1, int(request.GET.get("limit", 500))))
    rows = [
        {
            "id": r.id,
            "row_number": r.row_number,
            "source_submission_id": r.source_submission_id,
            "excluded": r.excluded,
            "data": r.data,
        }
        for r in run.records.order_by("row_number")[:limit]
    ]
    changes = {}
    for change in run.changes.order_by("id")[:10000]:
        key = f"{change.row_number}:{change.field}"
        changes.setdefault(key, []).append({
            "type": change.change_type,
            "old": change.old_value,
            "new": change.new_value,
            "detail": change.detail,
            "step": change.step.name if change.step_id else None,
        })
    return JsonResponse({
        "ok": True,
        "run": {
            "id": run.id,
            "status": run.status,
            "result_count": run.result_count,
            "column_count": run.column_count,
            "summary": run.summary,
        },
        "rows": rows,
        "changes": changes,
    })


@login_required
@require_POST
def run_dashboard(request, code, run_id):
    survey = _own(request, code)
    run = get_object_or_404(CleaningRun, id=run_id, pipeline__survey=survey, status="complete")
    data = _json(request) or {}
    dependent = data.get("dependent")
    independent = data.get("independent") or []
    return JsonResponse({
        "ok": True,
        "dashboard": dashboard_payload(
            run, dependent, independent, schema=_survey_schema(survey),
        ),
    })


@login_required
@require_GET
def run_summary(request, code, run_id):
    """Friendly per-column dataset summary."""
    survey = _own(request, code)
    run = get_object_or_404(CleaningRun, id=run_id, pipeline__survey=survey, status="complete")
    from .analytics import column_profile
    return JsonResponse({"ok": True, "summary": column_profile(run)})


@login_required
@require_POST
def run_chart(request, code, run_id):
    """Build one user-configured chart from a completed run."""
    survey = _own(request, code)
    run = get_object_or_404(CleaningRun, id=run_id, pipeline__survey=survey, status="complete")
    spec = _json(request) or {}
    from .analytics import custom_chart
    return JsonResponse(custom_chart(run, spec))


@login_required
@require_POST
def run_timeseries(request, code, run_id):
    """Aggregate a measure over time (with optional rolling average / grouping)."""
    survey = _own(request, code)
    run = get_object_or_404(CleaningRun, id=run_id, pipeline__survey=survey, status="complete")
    spec = _json(request) or {}
    from .analytics import time_series
    return JsonResponse(time_series(run, spec))


@login_required
@require_POST
def run_timeline(request, code, run_id):
    """Cumulative frames for the historical playback scrubber."""
    survey = _own(request, code)
    run = get_object_or_404(CleaningRun, id=run_id, pipeline__survey=survey, status="complete")
    spec = _json(request) or {}
    from .analytics import timeline_frames
    return JsonResponse(timeline_frames(run, spec, schema=_survey_schema(survey)))


@login_required
def run_board(request, code, run_id):
    """Load (GET) or save (POST) the customizable dashboard board for a run.

    The board definition is a JSON list of cards —
    {kind:"chart"|"kpi"|"map", spec, title, size} — stored in the
    AnalysisDashboard model so a curated dashboard survives reloads and
    can be reopened by anyone with access to the survey.
    """
    survey = _own(request, code)
    run = get_object_or_404(CleaningRun, id=run_id, pipeline__survey=survey, status="complete")

    if request.method == "POST":
        data = _json(request)
        if data is None or not isinstance(data.get("items"), list):
            return JsonResponse({"ok": False, "error": "items must be a list."}, status=400)
        board, _created = AnalysisDashboard.objects.get_or_create(
            run=run, name="Main board", defaults={"created_by": request.user},
        )
        board.definition = {"items": data["items"][:40]}
        board.save(update_fields=["definition", "updated_at"])
        return JsonResponse({"ok": True, "items": board.definition["items"]})

    board = run.dashboards.filter(name="Main board").order_by("-updated_at").first()
    return JsonResponse({
        "ok": True,
        "items": (board.definition or {}).get("items", []) if board else [],
    })


def _repeat_cell(value):
    """Human display for one cell in the clean-data Word export."""
    if isinstance(value, list):
        if value and isinstance(value[0], dict):
            return f"{len(value)} item(s)"
        return ", ".join(str(v) for v in value)
    if isinstance(value, dict):
        return f"{len(value)} field(s)"
    return "" if value is None else str(value)


def export_clean_word(run):
    """Clean-data-only Word document (no report machinery) — the third
    plain download format next to CSV and Excel."""
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed on the server.") from exc

    records = list(run.records.filter(excluded=False).order_by("row_number")[:5000])
    columns = [c["name"] for c in (run.schema or [])]
    if not columns:
        seen = {}
        for r in records:
            for k in (r.data or {}).keys():
                seen.setdefault(k, True)
        columns = list(seen.keys())

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    survey = run.pipeline.survey
    doc.add_heading(f"{survey.title} — clean dataset", level=1)
    meta = doc.add_paragraph(
        f"Pipeline: {run.pipeline.name} · Run #{run.id}"
        f" · {run.result_count} rows · {run.column_count} columns"
        + (f" · completed {run.completed_at:%d %b %Y %H:%M}" if run.completed_at else "")
    )
    meta.runs[0].font.size = Pt(9)
    if len(records) < run.result_count:
        doc.add_paragraph(
            f"Showing the first {len(records)} of {run.result_count} rows — "
            "use the CSV or Excel download for the complete dataset."
        ).runs[0].font.size = Pt(9)

    table = doc.add_table(rows=1, cols=len(columns) + 1)
    table.style = "Light Grid Accent 1"
    head = table.rows[0].cells
    head[0].text = "#"
    for i, col in enumerate(columns):
        head[i + 1].text = str(col)
    for r in records:
        cells = table.add_row().cells
        cells[0].text = str(r.row_number)
        data = r.data or {}
        for i, col in enumerate(columns):
            cells[i + 1].text = _repeat_cell(data.get(col))
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for rn in p.runs:
                    rn.font.size = Pt(8)

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = (
        f'attachment; filename="{survey.code}-run{run.id}-clean-data.docx"'
    )
    return response


@login_required
@require_GET
def run_export(request, code, run_id, fmt):
    survey = _own(request, code)
    run = get_object_or_404(CleaningRun, id=run_id, pipeline__survey=survey, status="complete")
    exporters = {
        # Full report exports (steps, stats, audit trail)
        "xlsx": export_excel,
        "docx": export_word,
        "pdf": export_pdf,
        "html": export_html,
        # Clean-data-only downloads
        "csv": export_clean_csv,
        "data-xlsx": export_clean_excel,
        "data-docx": export_clean_word,
    }
    exporter = exporters.get(fmt)
    if exporter is None:
        return JsonResponse({"ok": False, "error": "Unsupported format."}, status=400)
    try:
        return exporter(run)
    except RuntimeError as exc:
        return JsonResponse({"ok": False, "error": str(exc)}, status=501)