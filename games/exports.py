"""
Excel + Word export builders for a single LiveSession.

Each builder takes a LiveSession and returns a BytesIO containing the finished
file. The view layer wraps that in an HttpResponse with the right
Content-Disposition header.

Dependencies:
  - openpyxl       (Excel + native charts)
  - python-docx    (Word docs)
  - matplotlib     (chart images embedded in Word; openpyxl handles its own
                    charts natively so we don't need it for xlsx)

If matplotlib isn't installed the Word export still works — charts are just
omitted with a small "(chart unavailable)" placeholder line.

Fixed here
----------
* `doc.add_paragraph(...).italic = True` set an attribute on the Paragraph
  object, which has no such property. The text was never italic; the line
  just quietly did nothing. Formatting belongs on the run.
* Every table asked for the "Light Grid Accent 4" style by name. python-docx
  raises KeyError for any style missing from the template, so a project with
  a customised default.docx produced a 500 on export rather than a document.
  `_styled_table` degrades to whatever is available.
* The per-question chart was built as `type="bar"` (horizontal) while its
  axis titles described a vertical chart, so "Count" labelled the answers
  axis and "Answer" labelled the counts axis.
"""
from __future__ import annotations

from io import BytesIO

from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.chart.label import DataLabelList
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    import matplotlib
    matplotlib.use("Agg")  # no GUI; safe in a Django worker
    import matplotlib.pyplot as plt
    _HAS_MPL = True
except Exception:  # ImportError, or a broken backend on a headless box
    _HAS_MPL = False

from .leaderboards import (
    full_leaderboard,
    per_question_stats,
    room_standings,
    top_three_for_session,
)
from .models import GameAnswer

BRAND_PURPLE = "7C3AED"
BRAND_INK = "1E293B"


# ─────────────────────────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────────────────────────

def _safe_sheet_title(title: str) -> str:
    """Excel sheet names: no []:/\\?*, no leading/trailing apostrophe, ≤31 chars."""
    bad = '[]:/\\?*'
    cleaned = "".join("_" if ch in bad else ch for ch in title).strip().strip("'")
    return (cleaned or "Sheet")[:31].strip() or "Sheet"


def _unique_sheet_title(title: str, used: set) -> str:
    title = _safe_sheet_title(title)
    if title not in used:
        used.add(title)
        return title
    n = 2
    while True:
        suffix = f" ({n})"
        candidate = _safe_sheet_title(title[: 31 - len(suffix)] + suffix)
        if candidate not in used:
            used.add(candidate)
            return candidate
        n += 1


def _format_session_label(session) -> str:
    """Human-readable label used in headers and filenames."""
    when = session.created_at.strftime("%d %b %Y, %H:%M") if session.created_at else "unknown date"
    quiz_title = session.quiz.title if session.quiz else "Untitled"
    return f"{quiz_title} · session {session.code} · {when}"


def _fmt_dt(value):
    return value.strftime("%Y-%m-%d %H:%M:%S") if value else "—"


def _seconds(ms):
    return round((ms or 0) / 1000.0, 2)


# ─────────────────────────────────────────────────────────────────
# Excel export
# ─────────────────────────────────────────────────────────────────

_HEADER_FILL = PatternFill("solid", start_color=f"FF{BRAND_PURPLE}")
_HEADER_FONT = Font(name="Calibri", bold=True, color="FFFFFFFF", size=11)
_TITLE_FONT = Font(name="Calibri", bold=True, size=16, color=f"FF{BRAND_INK}")
_LABEL_FONT = Font(name="Calibri", bold=True, size=11)
_BODY_FONT = Font(name="Calibri", size=11)
_CORRECT_FILL = PatternFill("solid", start_color="FFD1FAE5")   # mint
_PODIUM_FILL = PatternFill("solid", start_color="FFFEF3C7")    # amber
_CENTER = Alignment(horizontal="center", vertical="center")


def _style_header_row(sheet, row, columns):
    for col_idx, _ in enumerate(columns, start=1):
        cell = sheet.cell(row=row, column=col_idx)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = _CENTER


def _autosize(sheet, max_width=48):
    """Best-effort column auto-width based on cell value lengths."""
    for col in sheet.columns:
        col_cells = list(col)
        if not col_cells:
            continue
        letter = get_column_letter(col_cells[0].column)
        longest = 0
        for cell in col_cells:
            if cell.value is None:
                continue
            longest = max(longest, len(str(cell.value)))
        sheet.column_dimensions[letter].width = min(max_width, max(10, longest + 2))


def build_excel(session) -> BytesIO:
    """Build a multi-sheet .xlsx workbook for one session and return BytesIO."""
    wb = Workbook()
    leaderboard = full_leaderboard(session)
    podium = top_three_for_session(session)
    qstats = per_question_stats(session)
    rooms = room_standings(session)

    # ── Overview ──
    ws = wb.active
    ws.title = "Overview"
    ws.sheet_properties.tabColor = BRAND_PURPLE

    ws["A1"] = "Knock-Knock game results"
    ws["A1"].font = _TITLE_FONT
    ws.merge_cells("A1:D1")

    facts = [
        ("Quiz", session.quiz.title if session.quiz else "—"),
        ("Session code", session.code),
        ("Started", _fmt_dt(session.created_at)),
        ("Ended", _fmt_dt(getattr(session, "ended_at", None))),
        ("State", session.get_state_display()),
        ("Players", len(leaderboard)),
        ("Questions", len(qstats)),
        ("Top scorer", f"{podium[0]['emoji']} {podium[0]['nickname']} — {podium[0]['score']} pts" if podium else "—"),
    ]
    for i, (label, value) in enumerate(facts, start=3):
        ws.cell(row=i, column=1, value=label).font = _LABEL_FONT
        ws.cell(row=i, column=2, value=value).font = _BODY_FONT
    _autosize(ws)

    # ── Leaderboard ──
    lb = wb.create_sheet("Leaderboard")
    headers = ["Rank", "Avatar", "Nickname", "Room", "Score"]
    lb.append(headers)
    _style_header_row(lb, 1, headers)
    for row in leaderboard:
        lb.append([row["rank"], row["emoji"], row["nickname"], row["room_id"], row["score"]])
    for i, row in enumerate(leaderboard, start=2):
        if row["rank"] <= 3:
            for col in range(1, len(headers) + 1):
                lb.cell(row=i, column=col).fill = _PODIUM_FILL
    _autosize(lb)
    lb.freeze_panes = "A2"

    # ── Rooms (only when the quiz uses them) ──
    if rooms:
        rs = wb.create_sheet("Rooms")
        room_headers = ["Rank", "Room", "Players", "Total points", "Average"]
        rs.append(room_headers)
        _style_header_row(rs, 1, room_headers)
        for r in rooms:
            rs.append([r["rank"], r["name"], r["members"], r["total"], r["average"]])
        _autosize(rs)
        rs.freeze_panes = "A2"

    # ── Per-question summary ──
    stats_ws = wb.create_sheet("Question stats")
    stats_headers = ["#", "Question", "Type", "Answers", "Correct", "Correct %", "Late", "Avg time (s)"]
    stats_ws.append(stats_headers)
    _style_header_row(stats_ws, 1, stats_headers)
    for idx, s in enumerate(qstats, start=1):
        q = s["question"]
        stats_ws.append([
            idx,
            q.text,
            q.get_question_type_display(),
            s["answer_count"],
            s["correct_count"],
            round(s["correct_pct"], 1),
            s["late_count"],
            _seconds(s["avg_time_ms"]),
        ])
    for r in range(2, len(qstats) + 2):
        stats_ws.cell(row=r, column=6).number_format = '0.0"%"'
    _autosize(stats_ws)
    stats_ws.freeze_panes = "A2"

    # ── Raw answers: one row per submission, for anyone who wants a pivot ──
    raw = wb.create_sheet("Raw answers")
    raw_headers = ["Question #", "Question", "Nickname", "Room", "Answer", "Correct", "Late", "Time (s)", "Points", "Submitted"]
    raw.append(raw_headers)
    _style_header_row(raw, 1, raw_headers)

    question_index = {s["question"].id: i for i, s in enumerate(qstats, start=1)}
    answers = (
        GameAnswer.objects
        .filter(session=session)
        .select_related("question", "choice")
        .order_by("question__order", "created_at")
    )
    for a in answers:
        if a.question.is_puzzle:
            answer_text = " → ".join(str(x) for x in (a.puzzle_order or [])) or "—"
        else:
            answer_text = a.choice.display_label if a.choice else "—"
        raw.append([
            question_index.get(a.question_id, ""),
            a.question.text,
            a.nickname,
            a.room_id,
            answer_text,
            "Yes" if a.is_correct else "No",
            "Yes" if a.was_late else "",
            _seconds(a.time_taken_ms),
            a.points_awarded,
            _fmt_dt(a.created_at),
        ])
    _autosize(raw, max_width=42)
    raw.freeze_panes = "A2"

    # ── One sheet per question with distribution + chart ──
    used_titles = {s.title for s in wb.worksheets}
    for idx, s in enumerate(qstats, start=1):
        q = s["question"]
        qs = wb.create_sheet(_unique_sheet_title(f"Q{idx} {q.text}", used_titles))

        qs["A1"] = f"Question {idx}: {q.text}"
        qs["A1"].font = _TITLE_FONT
        qs.merge_cells("A1:D1")

        summary = [
            ("Answers", s["answer_count"]),
            ("Correct", s["correct_count"]),
            ("Correct %", round(s["correct_pct"], 1)),
            ("Late", s["late_count"]),
            ("Avg time (s)", _seconds(s["avg_time_ms"])),
        ]
        for i, (label, value) in enumerate(summary, start=3):
            qs.cell(row=i, column=1, value=label).font = _LABEL_FONT
            qs.cell(row=i, column=2, value=value).font = _BODY_FONT
        qs.cell(row=5, column=2).number_format = '0.0"%"'

        header_row = 9
        dist_headers = ["Answer", "Count", "Correct?"]
        for col, header in enumerate(dist_headers, start=1):
            qs.cell(row=header_row, column=col, value=header)
        _style_header_row(qs, header_row, dist_headers)

        for i, (label, count, is_correct) in enumerate(s["distribution"], start=header_row + 1):
            qs.cell(row=i, column=1, value=label)
            qs.cell(row=i, column=2, value=count)
            qs.cell(row=i, column=3, value="✓" if is_correct else "")
            if is_correct:
                for c in range(1, 4):
                    qs.cell(row=i, column=c).fill = _CORRECT_FILL

        n = len(s["distribution"])
        if n:
            chart = BarChart()
            # "col" is vertical bars: answers along the bottom, counts up the
            # side, which is what the axis titles below actually describe.
            chart.type = "col"
            chart.style = 11
            chart.title = "Answer distribution"
            chart.x_axis.title = "Answer"
            chart.y_axis.title = "Responses"
            chart.legend = None
            chart.add_data(
                Reference(qs, min_col=2, min_row=header_row, max_row=header_row + n),
                titles_from_data=True,
            )
            chart.set_categories(
                Reference(qs, min_col=1, min_row=header_row + 1, max_row=header_row + n)
            )
            chart.dataLabels = DataLabelList(showVal=True)
            chart.height = 9   # cm
            chart.width = 16   # cm
            qs.add_chart(chart, "E3")

        _autosize(qs, max_width=60)

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────
# Word export
# ─────────────────────────────────────────────────────────────────

#: Tried in order. A project that ships a custom default.docx may have none
#: of the Accent variants, so plain "Table Grid" is the last resort and
#: `None` means "leave the built-in default alone".
_TABLE_STYLES = ("Light Grid Accent 4", "Light Grid Accent 1", "Table Grid")


def _styled_table(doc, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    for name in _TABLE_STYLES:
        try:
            table.style = name
            break
        except KeyError:
            continue
    return table


def _shade(cell, hex_fill):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), hex_fill))
    )


def _add_table_header(table, headers):
    hdr = table.rows[0].cells
    for i, txt in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        _shade(hdr[i], BRAND_PURPLE)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _label_value_row(table, row_idx, label, value):
    """Write a bold label / plain value pair without touching `runs[0]`.

    Setting `cell.text` then reaching for `paragraphs[0].runs[0]` raises
    IndexError the moment a value is an empty string, which is exactly what
    an unfinished session produces.
    """
    label_cell, value_cell = table.cell(row_idx, 0), table.cell(row_idx, 1)
    label_cell.text = ""
    label_cell.paragraphs[0].add_run(label).bold = True
    value_cell.text = str(value)


def _muted(paragraph, size=None):
    for run in paragraph.runs:
        run.italic = True
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        if size:
            run.font.size = Pt(size)
    return paragraph


def _chart_png(labels, values, correct_flags) -> BytesIO | None:
    """Render a horizontal bar chart with matplotlib and return a PNG BytesIO.

    Correct answers are coloured green, incorrect grey. Returns None if
    matplotlib is unavailable so callers can substitute a placeholder.
    """
    if not _HAS_MPL or not labels:
        return None

    colors = ["#10b981" if ok else "#94a3b8" for ok in correct_flags]
    fig, ax = plt.subplots(figsize=(6.2, max(2.2, 0.55 * len(labels) + 1)))
    bars = ax.barh(labels, values, color=colors, edgecolor="white")
    ax.invert_yaxis()
    ax.set_xlabel("Responses")

    # An all-zero question (nobody answered) made the old offset 0, stacking
    # every label on the axis line.
    peak = max(values) if values else 0
    offset = (peak * 0.02) if peak else 0.05
    if not peak:
        ax.set_xlim(0, 1)
    for bar, v in zip(bars, values):
        ax.text(bar.get_width() + offset, bar.get_y() + bar.get_height() / 2,
                str(v), va="center", fontsize=10)

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()

    out = BytesIO()
    fig.savefig(out, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    out.seek(0)
    return out


def build_word(session) -> BytesIO:
    """Build a .docx report for one session and return BytesIO."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ── Cover ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Knock-Knock game results")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(_format_session_label(session))
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph()

    # ── Session metadata ──
    meta = _styled_table(doc, rows=4, cols=2)
    _label_value_row(meta, 0, "Session code", session.code)
    _label_value_row(meta, 1, "Started", _fmt_dt(session.created_at))
    _label_value_row(meta, 2, "Ended", _fmt_dt(getattr(session, "ended_at", None)))
    _label_value_row(meta, 3, "State", session.get_state_display())

    doc.add_paragraph()

    # ── Podium ──
    podium = top_three_for_session(session)
    if podium:
        doc.add_heading("Top scorers", level=1)
        medals = ["🥇", "🥈", "🥉"]
        for i, p in enumerate(podium):
            para = doc.add_paragraph()
            r = para.add_run(f"{medals[i]}  {p['emoji']}  {p['nickname']}  —  {p['score']} pts")
            r.font.size = Pt(14)
            if i == 0:
                r.font.bold = True

    # ── Full leaderboard ──
    leaderboard = full_leaderboard(session)
    if leaderboard:
        doc.add_heading("Full leaderboard", level=1)
        table = _styled_table(doc, rows=1, cols=5)
        _add_table_header(table, ["Rank", "Avatar", "Nickname", "Room", "Score"])
        for p in leaderboard:
            row = table.add_row().cells
            row[0].text = str(p["rank"])
            row[1].text = p["emoji"]
            row[2].text = p["nickname"]
            row[3].text = p["room_id"]
            row[4].text = str(p["score"])
    else:
        _muted(doc.add_paragraph("Nobody joined this session."))

    # ── Room standings ──
    rooms = room_standings(session)
    if rooms:
        doc.add_heading("Room standings", level=1)
        _muted(doc.add_paragraph("Rooms are ranked by average score, so a small room isn't penalised for its size."))
        table = _styled_table(doc, rows=1, cols=5)
        _add_table_header(table, ["Rank", "Room", "Players", "Total", "Average"])
        for r in rooms:
            row = table.add_row().cells
            row[0].text = str(r["rank"])
            row[1].text = r["name"]
            row[2].text = str(r["members"])
            row[3].text = str(r["total"])
            row[4].text = str(r["average"])

    # ── Per-question sections ──
    qstats = per_question_stats(session)
    if qstats:
        doc.add_page_break()
        doc.add_heading("Question breakdown", level=1)

        for idx, s in enumerate(qstats, start=1):
            q = s["question"]
            doc.add_heading(f"Q{idx}. {q.text}", level=2)

            bits = [
                q.get_question_type_display(),
                f"{s['answer_count']} answers",
                f"{s['correct_count']} correct ({s['correct_pct']:.1f}%)",
                f"avg {_seconds(s['avg_time_ms'])}s",
            ]
            if s["late_count"]:
                bits.append(f"{s['late_count']} late")
            facts = doc.add_paragraph()
            facts.add_run("  ·  ".join(bits))
            _muted(facts)

            dist = s["distribution"]
            if not dist:
                continue

            table = _styled_table(doc, rows=1, cols=3)
            _add_table_header(table, ["Answer", "Count", "Correct?"])
            for label, count, is_correct in dist:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = str(count)
                row[2].text = "✓" if is_correct else ""
                if is_correct:
                    for cell in row:
                        _shade(cell, "D1FAE5")

            png = _chart_png([d[0] for d in dist], [d[1] for d in dist], [d[2] for d in dist])
            if png is not None:
                doc.add_paragraph()
                doc.add_picture(png, width=Inches(6.0))
            else:
                _muted(doc.add_paragraph("Chart unavailable — matplotlib is not installed on the server."))

            doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
